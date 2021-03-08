from docx import Document
from docx.shared import Pt
from docx.shared import Mm
from docx.shared import Cm
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from time import strftime
from time import localtime
from datetime import datetime
from docx.enum.text import WD_LINE_SPACING
from math import ceil
import os
import pandas as pd
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

print('外挂支撑8工况应力云图自动提取到word强化版--6杆轴力版本')
print('[可读取TheResult.txt结果]')
print('20210305 by 徐明')
print('')
'''
TODO:
1. 将图片结果提取设置成可选项
2. 计算构件名称可选填（外挂架，支撑系统，底座梁等等）
3. 单独组件应力图数量可选，组件名称可选填
4. 节点反力数量可选，增加包络值可选
5. 杆件轴力数量可选
'''
calc_book = Document()
# 设置正文字体
calc_book.styles['Normal'].font.name = 'Times New Roman'
calc_book.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
calc_book.styles['Normal'].font.size = Pt(12)
calc_book.styles['Normal'].font.color.rgb = RGBColor(0x00, 0x00, 0x00)
calc_book.styles['Normal'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
# 首行缩进，字符宽度等于字符高度，12pt=4.23mm, 1pt=0.3527mm
calc_book.styles['Normal'].paragraph_format.first_line_indent = Mm(8.46)
calc_book.styles['Normal'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
calc_book.styles['Normal'].paragraph_format.space_before = Mm(0)
calc_book.styles['Normal'].paragraph_format.space_after = Mm(0)
# 增加一个居中显示的样式
calc_book.styles['No Spacing'].font.name = 'Times New Roman'
calc_book.styles['No Spacing'].element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
calc_book.styles['No Spacing'].font.size = Pt(12)
calc_book.styles['No Spacing'].font.color.rgb = RGBColor(0x00, 0x00, 0x00)
calc_book.styles['No Spacing'].paragraph_format.first_line_indent = Mm(0)
calc_book.styles['No Spacing'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
calc_book.styles['No Spacing'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
calc_book.styles['No Spacing'].paragraph_format.space_before = Mm(0)
calc_book.styles['No Spacing'].paragraph_format.space_after = Mm(0)
# 设置标题1字体
calc_book.styles['Heading 1'].font.name = 'Times New Roman'
calc_book.styles['Heading 1'].element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
calc_book.styles['Heading 1'].element.rPr.rFonts.set(qn('w:asciiTheme'), 'Times New Roman')
calc_book.styles['Heading 1'].element.rPr.rFonts.set(qn('w:eastAsiaTheme'), '微软雅黑')
calc_book.styles['Heading 1'].element.rPr.rFonts.set(qn('w:hAnsiTheme'), 'Times New Roman')
calc_book.styles['Heading 1'].element.rPr.rFonts.set(qn('w:cstheme'), 'Times New Roman')
calc_book.styles['Heading 1'].font.size = Pt(16)
calc_book.styles['Heading 1'].font.color.rgb = RGBColor(0x00, 0x00, 0x00)
calc_book.styles['Heading 1'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
calc_book.styles['Heading 1'].paragraph_format.first_line_indent = Mm(0)
calc_book.styles['Heading 1'].paragraph_format.keep_with_next = True
calc_book.styles['Heading 1'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
calc_book.styles['Heading 1'].paragraph_format.space_before = Mm(0)
calc_book.styles['Heading 1'].paragraph_format.space_after = Mm(0)
# 设置标题2字体
calc_book.styles['Heading 2'].font.name = 'Times New Roman'
calc_book.styles['Heading 2'].element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
calc_book.styles['Heading 2'].element.rPr.rFonts.set(qn('w:asciiTheme'), 'Times New Roman')
calc_book.styles['Heading 2'].element.rPr.rFonts.set(qn('w:eastAsiaTheme'), '微软雅黑')
calc_book.styles['Heading 2'].element.rPr.rFonts.set(qn('w:hAnsiTheme'), 'Times New Roman')
calc_book.styles['Heading 2'].element.rPr.rFonts.set(qn('w:cstheme'), 'Times New Roman')
calc_book.styles['Heading 2'].font.size = Pt(14)
calc_book.styles['Heading 2'].font.color.rgb = RGBColor(0x00, 0x00, 0x00)
calc_book.styles['Heading 2'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
calc_book.styles['Heading 2'].paragraph_format.first_line_indent = Mm(0)
calc_book.styles['Heading 2'].paragraph_format.keep_with_next = True
calc_book.styles['Heading 2'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
calc_book.styles['Heading 2'].paragraph_format.space_before = Mm(0)
calc_book.styles['Heading 2'].paragraph_format.space_after = Mm(0)
# 设置标题3字体
calc_book.styles['Heading 3'].font.name = 'Times New Roman'
calc_book.styles['Heading 3'].element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
calc_book.styles['Heading 3'].element.rPr.rFonts.set(qn('w:asciiTheme'), 'Times New Roman')
calc_book.styles['Heading 3'].element.rPr.rFonts.set(qn('w:eastAsiaTheme'), '微软雅黑')
calc_book.styles['Heading 3'].element.rPr.rFonts.set(qn('w:hAnsiTheme'), 'Times New Roman')
calc_book.styles['Heading 3'].element.rPr.rFonts.set(qn('w:cstheme'), 'Times New Roman')
calc_book.styles['Heading 3'].font.size = Pt(14)
calc_book.styles['Heading 3'].font.color.rgb = RGBColor(0x00, 0x00, 0x00)
calc_book.styles['Heading 3'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
calc_book.styles['Heading 3'].paragraph_format.first_line_indent = Mm(0)
calc_book.styles['Heading 3'].paragraph_format.keep_with_next = True
calc_book.styles['Heading 3'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
calc_book.styles['Heading 3'].paragraph_format.space_before = Mm(0)
calc_book.styles['Heading 3'].paragraph_format.space_after = Mm(0)

'''
计算书文档属性设置
'''
print('文档属性初始化……')
# 文档标题
calc_book.core_properties.title = '计算书'
# 文档主题
calc_book.core_properties.subject = '计算书'
# 文档作者
calc_book.core_properties.author = 'xuming'
# 文档类别
calc_book.core_properties.category = 'calculation sheet'
# 文档注释
calc_book.core_properties.comments = 'Designed By Xuming. All Rights Reserved.'
# 文档创建时间
calc_book.core_properties.created = datetime.utcnow()
# 文档修改时间
calc_book.core_properties.modified = datetime.utcnow()

# 准确定义工作目录为文件所在目录
os.chdir(os.path.dirname(__file__))
'''
计算书正文开始
'''
# calc_book.add_heading('四、外挂支撑系统有限元建模分析', level=1)
# calc_book.add_paragraph('采用大型通用有限元软件ANSYS按“三、外挂支撑系统图纸”建模，按“二、结构反力”加载分析', style='Normal')
# calc_book.add_paragraph('取塔机吊臂旋转每隔45度为一个工况，分析8种工况。', style='Normal')

titlist = input('输入上级二级标题编号[例如4.1，4.2，4.3等]: ')
# titlist = '4.4'
while True:
    try:
        jobname = int(input('输入Job Name工作名[例如：850，1250，1500等]：'))
        break
    except ValueError:
        print("输入错误，请输入正确的Job Name工作名[整数]")

picnum = int(jobname * 1000)
skiplist = [0, 1, 2, 27, 28, 53, 54, 79, 80, 105, 106, 131, 132, 157, 158, 183, 184,
            209, 210, 211, 236, 237, 262, 263, 288]

ansys = pd.read_table('TheResult.txt', sep='|', skiprows=skiplist, header=None)
# print(ansys)
tabindex = 12

# calc_book.add_heading(f'{titlist}.ZSL850外挂支撑系统', level=2)
# calc_book.add_paragraph('', style='Normal')
# calc_book.add_paragraph('外挂架模型', style='No Spacing')
# calc_book.add_paragraph('', style='Normal')
# calc_book.add_paragraph('8种工况示意图', style='No Spacing')

for i in range(1, 9):
    calc_book.add_heading(f'{titlist}.{i}.工况{i}', level=3)
    max1 = ansys.iloc[tabindex, 1]
    max1 = abs(max1)
    min1 = ansys.iloc[tabindex + 1, 1]
    min1 = abs(min1)
    outnum1 = max(max1, min1)

    max2 = ansys.iloc[tabindex + 2, 1]
    max2 = abs(max2)
    min2 = ansys.iloc[tabindex + 3, 1]
    min2 = abs(min2)
    outnum2 = max(max2, min2)

    max3 = ansys.iloc[tabindex + 4, 1]
    max3 = abs(max3)
    min3 = ansys.iloc[tabindex + 5, 1]
    min3 = abs(min3)
    outnum3 = max(max3, min3)
    calc_book.add_paragraph(f'外挂架X向最大挠度{ceil(outnum1)}mm，Y向最大挠度{ceil(outnum2)}mm，Z向最大挠度{ceil(outnum3)}mm:',
                            style='Normal')
    calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{picnum}.png', height=Cm(7))
    picnum = picnum + 1
    calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{picnum}.png', height=Cm(7))
    picnum = picnum + 1
    calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{picnum}.png', height=Cm(7))
    picnum = picnum + 1
    max4 = ansys.iloc[tabindex + 6, 1]
    max4 = abs(max4)
    min4 = ansys.iloc[tabindex + 7, 1]
    min4 = abs(min4)
    outnum4 = max(max4, min4)
    calc_book.add_paragraph(f'主梁最大Mises应力{ceil(outnum4)}MPa:', style='Normal')
    calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{picnum}.png', height=Cm(7))
    picnum = picnum + 1
    max5 = ansys.iloc[tabindex + 8, 1]
    max5 = abs(max5)
    min5 = ansys.iloc[tabindex + 9, 1]
    min5 = abs(min5)
    outnum5 = max(max5, min5)
    calc_book.add_paragraph(f'竖向撑杆最大Mises应力{ceil(outnum5)}MPa:', style='Normal')
    calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{picnum}.png', height=Cm(7))
    picnum = picnum + 1
    max6 = ansys.iloc[tabindex + 10, 1]
    max6 = abs(max6)
    min6 = ansys.iloc[tabindex + 11, 1]
    min6 = abs(min6)
    outnum6 = max(max6, min6)
    calc_book.add_paragraph(f'水平撑杆最大Mises应力{ceil(outnum6)}MPa:', style='Normal')
    calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{picnum}.png', height=Cm(7))
    picnum = picnum + 1
    tabindex = tabindex + 24  # 关键

calc_book.add_heading(f'{titlist}.9.节点反力汇总', level=3)

rows_num = 64
cols_num = 4

mytable = calc_book.add_table(rows=rows_num, cols=cols_num, style='Table Grid')
mytable.alignment = WD_TABLE_ALIGNMENT.CENTER
mytable.autofit = False



for r in range(8):
    mytable.cell(0 + r * 8, 0).text = f'工况-{r + 1}'
    mytable.cell(1 + r * 8, 0).text = '节点'
    mytable.cell(1 + r * 8, 1).text = 'F_x(t)'
    mytable.cell(1 + r * 8, 2).text = 'F_y(t)'
    mytable.cell(1 + r * 8, 3).text = 'F_z(t)'
    for i in range(6):
        mytable.cell(2 + i + r * 8, 0).text = str(int(ansys.iloc[i + r * 24, 0]))
        for j in range(1, 4):
            mytable.cell(2 + i + r * 8, j).text = str(ansys.iloc[i + r * 24, j])

# 设置单元格格式
for c in range(cols_num):
    for cell in mytable.columns[c].cells:
        cell.width = Cm(3.5)
        cell.paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        cell.paragraphs[0].paragraph_format.first_line_indent = Mm(0)
        cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
for row in mytable.rows:
    row.height = Cm(0.8)

calc_book.add_paragraph('', style='Normal')
calc_book.add_paragraph('杆件轴力包络值', style='Normal')
mytable2 = calc_book.add_table(rows=7, cols=2, style='Table Grid')
mytable2.alignment = WD_TABLE_ALIGNMENT.CENTER
mytable2.autofit = False
mytable2.cell(0, 0).text = '名称'
# mytable2.cell(0, 1).text = '杆件长度/mm'
mytable2.cell(0, 1).text = '轴向力/t'

mytable2.cell(1, 0).text = '水平杆1'

mytable2.cell(1, 1).text = str(ceil(ansys.iloc[6 + 10 * 24, 1]))

mytable2.cell(2, 0).text = '水平杆2'

mytable2.cell(2, 1).text = str(ceil(ansys.iloc[7 + 10 * 24, 1]))

mytable2.cell(3, 0).text = '竖向撑杆1'

mytable2.cell(3, 1).text = str(ceil(ansys.iloc[8 + 10 * 24, 1]))

mytable2.cell(4, 0).text = '竖向撑杆2'
mytable2.cell(4, 1).text = str(ceil(ansys.iloc[9 + 10 * 24, 1]))

mytable2.cell(5, 0).text = '竖向撑杆3'

mytable2.cell(5, 1).text = str(ceil(ansys.iloc[10 + 10 * 24, 1]))

mytable2.cell(6, 0).text = '竖向撑杆4'
mytable2.cell(6, 1).text = str(ceil(ansys.iloc[11 + 10 * 24, 1]))

# 设置单元格格式
for c in range(2):
    for cell in mytable2.columns[c].cells:
        cell.width = Cm(3.5)
        cell.paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        cell.paragraphs[0].paragraph_format.first_line_indent = Mm(0)
        cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
for row in mytable2.rows:
    row.height = Cm(0.8)

'''
计算书结束，输出docx文档
'''
filename = f'{jobname}外挂支撑系统计算结果' + strftime("%Y-%m-%d-%H%M%S", localtime())
calc_book.save(f'{filename}.docx')
print(f'计算书生成结束，保存在程序目录下，文件名为{filename}.docx')
# input('按回车键退出......')
