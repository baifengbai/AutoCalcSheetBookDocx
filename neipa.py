from docx import Document
# from docx.shared import Inches
from docx.shared import Pt
from docx.shared import Mm
from docx.shared import Cm
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
# import urllib.parse
# import urllib.request
# from PIL import Image
from time import strftime
from time import localtime
from datetime import datetime
from docx.enum.text import WD_LINE_SPACING
from math import ceil
# from os import mkdir
# from os.path import exists
import os
import pandas as pd
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

print('内爬支撑8工况应力云图自动提取到word强化版')
print('[可读取TheResult.txt结果]')
print('20210122 by 徐明')
print('')

calc_book = Document()
# 设置正文字体
calc_book.styles['Normal'].font.name = 'Times New Roman'
calc_book.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
calc_book.styles['Normal'].font.size = Pt(12)
calc_book.styles['Normal'].font.color.rgb = RGBColor(0x00, 0x00, 0x00)
calc_book.styles['Normal'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
# 首行缩进，字符宽度等于字符高度，12pt=4.23mm, 1pt=0.3527mm
calc_book.styles['Normal'].paragraph_format.first_line_indent = Mm(8.46)
calc_book.styles['Normal'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
calc_book.styles['Normal'].paragraph_format.space_before = Mm(0)
calc_book.styles['Normal'].paragraph_format.space_after = Mm(0)
# 增加一个居中显示的样式
calc_book.styles['No Spacing'].font.name = 'Times New Roman'
calc_book.styles['No Spacing'].element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
calc_book.styles['No Spacing'].font.size = Pt(12)
calc_book.styles['No Spacing'].font.color.rgb = RGBColor(0x00, 0x00, 0x00)
calc_book.styles['No Spacing'].paragraph_format.first_line_indent = Mm(0)
calc_book.styles['No Spacing'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
calc_book.styles['No Spacing'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
calc_book.styles['No Spacing'].paragraph_format.space_before = Mm(0)
calc_book.styles['No Spacing'].paragraph_format.space_after = Mm(0)
# 设置标题1字体
calc_book.styles['Heading 1'].font.name = 'Times New Roman'
calc_book.styles['Heading 1'].element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
calc_book.styles['Heading 1'].element.rPr.rFonts.set(qn('w:asciiTheme'), 'Times New Roman')
calc_book.styles['Heading 1'].element.rPr.rFonts.set(qn('w:eastAsiaTheme'), '微软雅黑')
calc_book.styles['Heading 1'].element.rPr.rFonts.set(qn('w:hAnsiTheme'), 'Times New Roman')
calc_book.styles['Heading 1'].element.rPr.rFonts.set(qn('w:cstheme'), 'Times New Roman')
calc_book.styles['Heading 1'].font.size = Pt(16)
calc_book.styles['Heading 1'].font.color.rgb = RGBColor(0x00, 0x00, 0x00)
calc_book.styles['Heading 1'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
calc_book.styles['Heading 1'].paragraph_format.first_line_indent = Mm(0)
calc_book.styles['Heading 1'].paragraph_format.keep_with_next = True
calc_book.styles['Heading 1'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
calc_book.styles['Heading 1'].paragraph_format.space_before = Mm(0)
calc_book.styles['Heading 1'].paragraph_format.space_after = Mm(0)
# 设置标题2字体
calc_book.styles['Heading 2'].font.name = 'Times New Roman'
calc_book.styles['Heading 2'].element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
calc_book.styles['Heading 2'].element.rPr.rFonts.set(qn('w:asciiTheme'), 'Times New Roman')
calc_book.styles['Heading 2'].element.rPr.rFonts.set(qn('w:eastAsiaTheme'), '微软雅黑')
calc_book.styles['Heading 2'].element.rPr.rFonts.set(qn('w:hAnsiTheme'), 'Times New Roman')
calc_book.styles['Heading 2'].element.rPr.rFonts.set(qn('w:cstheme'), 'Times New Roman')
calc_book.styles['Heading 2'].font.size = Pt(14)
calc_book.styles['Heading 2'].font.color.rgb = RGBColor(0x00, 0x00, 0x00)
calc_book.styles['Heading 2'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
calc_book.styles['Heading 2'].paragraph_format.first_line_indent = Mm(0)
calc_book.styles['Heading 2'].paragraph_format.keep_with_next = True
calc_book.styles['Heading 2'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
calc_book.styles['Heading 2'].paragraph_format.space_before = Mm(0)
calc_book.styles['Heading 2'].paragraph_format.space_after = Mm(0)
# 设置标题3字体
calc_book.styles['Heading 3'].font.name = 'Times New Roman'
calc_book.styles['Heading 3'].element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
calc_book.styles['Heading 3'].element.rPr.rFonts.set(qn('w:asciiTheme'), 'Times New Roman')
calc_book.styles['Heading 3'].element.rPr.rFonts.set(qn('w:eastAsiaTheme'), '微软雅黑')
calc_book.styles['Heading 3'].element.rPr.rFonts.set(qn('w:hAnsiTheme'), 'Times New Roman')
calc_book.styles['Heading 3'].element.rPr.rFonts.set(qn('w:cstheme'), 'Times New Roman')
calc_book.styles['Heading 3'].font.size = Pt(14)
calc_book.styles['Heading 3'].font.color.rgb = RGBColor(0x00, 0x00, 0x00)
calc_book.styles['Heading 3'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
calc_book.styles['Heading 3'].paragraph_format.first_line_indent = Mm(0)
calc_book.styles['Heading 3'].paragraph_format.keep_with_next = True
calc_book.styles['Heading 3'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
calc_book.styles['Heading 3'].paragraph_format.space_before = Mm(0)
calc_book.styles['Heading 3'].paragraph_format.space_after = Mm(0)

'''
计算书文档属性设置
'''
print('文档属性初始化……')
# 文档标题
calc_book.core_properties.title = '计算书'
# 文档主题
calc_book.core_properties.subject = '计算书'
# 文档作者
calc_book.core_properties.author = 'xuming'
# 文档类别
calc_book.core_properties.category = 'calculation sheet'
# 文档注释
calc_book.core_properties.comments = 'Designed By Xuming. All Rights Reserved.'
# 文档创建时间
calc_book.core_properties.created = datetime.utcnow()
# 文档修改时间
calc_book.core_properties.modified = datetime.utcnow()

# 准确定义工作目录为文件所在目录
os.chdir(os.path.dirname(__file__))
'''
计算书正文开始
'''

titlist = input('输入上级二级标题编号[例如4.1，4.2，4.3等]: ')
# titlist = '4.4'
while True:
    try:
        jobname = int(input('输入Job Name工作名[例如：850，1250，1500等]：'))
        break
    except ValueError:
        print("输入错误，请输入正确的Job Name工作名[整数]")

picnum = int(jobname * 1000)
skiplist = [0, 1, 2, 15, 16, 29, 30, 43, 44, 57, 58, 71, 72, 85, 86, 99, 100,
            113, 114, 115, 128, 129, 142, 143, 156]

ansys = pd.read_table('TheResult.txt', sep='|', skiprows=skiplist, header=None)
# print(ansys)
tabindex = 4


for i in range(1, 9):
    calc_book.add_heading(f'{titlist}.{i}.工况{i}', level=3)
    max1 = ansys.iloc[tabindex, 1]
    max1 = abs(max1)
    min1 = ansys.iloc[tabindex + 1, 1]
    min1 = abs(min1)
    outnum1 = max(max1, min1)

    max2 = ansys.iloc[tabindex + 2, 1]
    max2 = abs(max2)
    min2 = ansys.iloc[tabindex + 3, 1]
    min2 = abs(min2)
    outnum2 = max(max2, min2)

    max3 = ansys.iloc[tabindex + 4, 1]
    max3 = abs(max3)
    min3 = ansys.iloc[tabindex + 5, 1]
    min3 = abs(min3)
    outnum3 = max(max3, min3)
    calc_book.add_paragraph(f'内爬钢梁X向最大挠度{ceil(outnum1)}mm，Y向最大挠度{ceil(outnum2)}mm，Z向最大挠度{ceil(outnum3)}mm:',
                            style='Normal')
    calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{picnum}.png', height=Cm(7))
    picnum = picnum + 1
    calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{picnum}.png', height=Cm(7))
    picnum = picnum + 1
    calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{picnum}.png', height=Cm(7))
    picnum = picnum + 1
    max4 = ansys.iloc[tabindex + 6, 1]
    max4 = abs(max4)
    min4 = ansys.iloc[tabindex + 7, 1]
    min4 = abs(min4)
    outnum4 = max(max4, min4)
    calc_book.add_paragraph(f'内爬钢梁最大Mises应力{ceil(outnum4)}MPa:', style='Normal')
    calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{picnum}.png', height=Cm(7))
    picnum = picnum + 1
    tabindex = tabindex + 12  # 关键

calc_book.add_heading(f'{titlist}.9.节点反力汇总', level=3)

rows_num = 48
cols_num = 4

mytable = calc_book.add_table(rows=rows_num, cols=cols_num, style='Table Grid')
mytable.alignment = WD_TABLE_ALIGNMENT.CENTER
mytable.autofit = False

for r in range(8):  # 8工况循环
    mytable.cell(0 + r * 6, 0).text = f'工况-{r + 1}'
    mytable.cell(1 + r * 6, 0).text = '节点'
    mytable.cell(1 + r * 6, 1).text = 'F_x(t)'
    mytable.cell(1 + r * 6, 2).text = 'F_y(t)'
    mytable.cell(1 + r * 6, 3).text = 'F_z(t)'
    for i in range(4):  # 4节点循环
        mytable.cell(2 + i + r * 6, 0).text = str(int(ansys.iloc[i + r * 12, 0]))
        for j in range(1, 4):   # 三向反力提取
            mytable.cell(2 + i + r * 6, j).text = str(ansys.iloc[i + r * 12, j])

# 设置单元格格式
for c in range(cols_num):
    for cell in mytable.columns[c].cells:
        cell.width = Cm(3.5)
        cell.paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        cell.paragraphs[0].paragraph_format.first_line_indent = Mm(0)
        cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
for row in mytable.rows:
    row.height = Cm(0.8)


'''
计算书结束，输出docx文档
'''
filename = f'{jobname}内爬支撑系统计算结果' + strftime("%Y-%m-%d-%H%M%S", localtime())
calc_book.save(f'{filename}.docx')
print(f'计算书生成结束，保存在程序目录下，文件名为{filename}.docx')
# input('按回车键退出......')
