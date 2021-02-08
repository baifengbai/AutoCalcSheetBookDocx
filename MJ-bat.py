from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import Mm
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PIL import Image, ImageChops
from time import strftime
from time import localtime
from datetime import datetime
from docx.enum.text import WD_LINE_SPACING
from math import ceil
import matplotlib.pyplot as plt
from matplotlib import rcParams
import os
import pandas as pd
import sys


# 准确定义工作目录为文件所在目录
os.chdir(os.path.dirname(__file__))

config = {
    "font.family": 'serif',
    "font.size": 20,
    "mathtext.fontset": 'stix',
    "font.serif": ['SimSun'],
}
rcParams.update(config)

print('======================================')
print('[批量]墙体埋件计算书word生成器')
print('测试版本20210129 by 徐明')
print('默认参数：fy=300；fc=19.1；ft=1.71')
print('======================================')
print('计算书开始生成……')

path = 'images'
if not os.path.exists(path):
    os.mkdir(path)


def trim(im2):
    bg = Image.new(im2.mode, im2.size, im2.getpixel((0, 0)))
    diff = ImageChops.difference(im2, bg)
    diff = ImageChops.add(diff, diff, 2.0, -100)
    bbox = diff.getbbox()
    if bbox:
        return im2.crop(bbox)


# 插入公式图片函数，参数1公式字符串，参数2公式图片的名称，返回值为公式图片的原始英寸宽度
def add_image(latex, jpgname):
    fig = plt.figure(figsize=(20, 10), dpi=300)
    ax = fig.add_axes([0, 0, 1, 1])
    ax.get_xaxis().set_visible(False)
    ax.get_yaxis().set_visible(False)
    ax.spines['top'].set_visible(False)
    ax.spines['bottom'].set_visible(False)
    ax.spines['left'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.set_xticks([])
    ax.set_yticks([])
    str_latex = '$' + latex + '$'
    plt.text(0.5, 0.5, str_latex, fontsize=4.5, verticalalignment='center', horizontalalignment='center')
    plt.axis('off')
    plt.savefig(f'{path}/{jpgname}.jpg')
    plt.cla()
    plt.close("all")
    im = Image.open(f'{path}/{jpgname}.jpg')
    im = trim(im)
    im.save(f'{path}/{jpgname}.jpg')
    return im.size[0] / 96


# 2021-01-29 改进版：从excel表单中读取数据，批量生成计算书
try:
    data = pd.read_excel('maijian.xlsx', sheet_name='Sheet1', usecols='A:M',
                     converters={'锚筋直径/mm': int, '锚筋牌号': str, '锚板厚度/mm': float, '锚板高度/mm': float,
                                 '锚板宽度/mm': float, '水平向锚筋间距/mm': float, '竖向锚筋间距/mm': float, '水平锚筋排数': int,
                                 '竖向锚筋排数': int, '水平向剪力/N': float, '竖向剪力/N': float, '法向拉(+)压(-)力/N': float,
                                 '剪力作用点距离锚板平面的距离/mm': float
                                 })
except:
    print('maijian.xlsx不在同目录下')
    sys.exit()

calc_num = data.shape[0]
for calc_i in range(0, calc_num):
    # 创建文档对象
    calc_book = Document()
    # 设置正文字体
    calc_book.styles['Normal'].font.name = 'Italic'
    calc_book.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    calc_book.styles['Normal'].font.size = Pt(12)
    calc_book.styles['Normal'].font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    calc_book.styles['Normal'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    # 首行缩进，字符宽度等于字符高度，12pt=4.23mm, 1pt=0.3527mm
    calc_book.styles['Normal'].paragraph_format.first_line_indent = Mm(8.46)
    calc_book.styles['Normal'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    calc_book.styles['Normal'].paragraph_format.space_before = Mm(0)
    calc_book.styles['Normal'].paragraph_format.space_after = Mm(0)
    # 增加一个居中显示的样式
    calc_book.styles['No Spacing'].paragraph_format.first_line_indent = Mm(0)
    calc_book.styles['No Spacing'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    calc_book.styles['No Spacing'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    calc_book.styles['No Spacing'].paragraph_format.space_before = Mm(0)
    calc_book.styles['No Spacing'].paragraph_format.space_after = Mm(0)
    # 增加一个右对齐的样式
    calc_book.styles['Quote'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    calc_book.styles['Quote'].font.name = 'Italic'
    calc_book.styles['Quote'].element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    calc_book.styles['Quote'].font.size = Pt(12)
    calc_book.styles['Quote'].font.italic = False
    calc_book.styles['Quote'].font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    calc_book.styles['Quote'].paragraph_format.space_before = Mm(0)
    calc_book.styles['Quote'].paragraph_format.space_after = Mm(0)
    calc_book.styles['Quote'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    calc_book.styles['Quote'].paragraph_format.first_line_indent = Mm(0)
    # 设置标题字体
    calc_book.styles['Heading 1'].font.name = 'Times New Roman'
    calc_book.styles['Heading 1'].element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    calc_book.styles['Heading 1'].element.rPr.rFonts.set(qn('w:asciiTheme'), 'Times New Roman')
    calc_book.styles['Heading 1'].element.rPr.rFonts.set(qn('w:eastAsiaTheme'), '微软雅黑')
    calc_book.styles['Heading 1'].element.rPr.rFonts.set(qn('w:hAnsiTheme'), 'Times New Roman')
    calc_book.styles['Heading 1'].element.rPr.rFonts.set(qn('w:cstheme'), 'Times New Roman')
    calc_book.styles['Heading 1'].font.size = Pt(22)
    calc_book.styles['Heading 1'].font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    calc_book.styles['Heading 1'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    calc_book.styles['Heading 1'].paragraph_format.first_line_indent = Mm(0)
    calc_book.styles['Heading 1'].paragraph_format.keep_with_next = True
    calc_book.styles['Heading 1'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    calc_book.styles['Heading 1'].paragraph_format.space_before = Mm(0)
    calc_book.styles['Heading 1'].paragraph_format.space_after = Mm(0)

    '''
    计算书文档属性设置
    '''
    print('文档属性初始化……')
    # 文档标题
    calc_book.core_properties.title = '计算书'
    # 文档主题
    calc_book.core_properties.subject = '计算书'
    # 文档作者
    calc_book.core_properties.author = 'xuming'
    # 文档类别
    calc_book.core_properties.category = 'calculation sheet'
    # 文档注释
    calc_book.core_properties.comments = 'Designed By Xuming. All Rights Reserved.'
    # 文档创建时间
    calc_book.core_properties.created = datetime.utcnow()
    # 文档修改时间
    calc_book.core_properties.modified = datetime.utcnow()
    print(f'开始生成第{calc_i + 1}个计算书')
    d = data.loc[calc_i, '锚筋直径/mm']
    gjph = data.loc[calc_i, '锚筋牌号']
    mj = f'{gjph}-ø{d}'  # 锚筋规格
    mbt = data.loc[calc_i, '锚板厚度/mm']
    mbh = data.loc[calc_i, '锚板高度/mm']
    mbw = data.loc[calc_i, '锚板宽度/mm']
    mjjjx = data.loc[calc_i, '水平向锚筋间距/mm']
    mjjjz = data.loc[calc_i, '竖向锚筋间距/mm']
    xrow = data.loc[calc_i, '水平锚筋排数']
    zrow = data.loc[calc_i, '竖向锚筋排数']
    vx = data.loc[calc_i, '水平向剪力/N']
    vz = data.loc[calc_i, '竖向剪力/N']
    ny = data.loc[calc_i, '法向拉(+)压(-)力/N']
    vl = data.loc[calc_i, '剪力作用点距离锚板平面的距离/mm']
    fy = 300  # 锚筋的抗拉强度设计值，单位N/mm^2
    fc = 19.1  # 混凝土轴心抗压强度设计值，单位N/mm^2
    ft = 1.71  # 混凝土轴心抗拉强度设计值，单位N/mm^2
    zx = mjjjx * (xrow - 1)  # x向剪力作用方向最外层锚筋中心线之间的距离，单位mm
    zz = mjjjz * (zrow - 1)  # z向剪力作用方向最外层锚筋中心线之间的距离，单位mm
    row = max(xrow, zrow)  # 锚筋排数取大值
    mjnums = xrow * zrow  # 锚筋数量
    mz = vx * vl  # z向弯矩设计值，由x向剪力产生，单位N.mm
    mx = vz * vl  # x向弯矩设计值，由z向剪力产生，单位N.mm
    if row == 2:  # 锚筋层数的影响系数；当锚筋按等间距布置时：两层取1.0；三层取0.9；四层取0.85
        ar = 1
    elif row == 3:
        ar = 0.9
    elif row == 4:
        ar = 0.85
    else:
        ar = 0.85
    av = (4 - 0.08 * d) * ((fc / fy) ** 0.5)  # 锚筋的受剪承载力系数
    if av > 0.7:  # 根据规范大于0.7时取0.7
        av = 0.7
    ab = 0.6 + 0.25 * mbt / d  # 锚板的弯曲变形折减系数
    mjas = mjnums * 3.14 * ((d / 2) ** 2)  # 锚筋的总面积
    calc_book.add_paragraph('计算依据：GB50010-2010混凝土设计规范(2015年版)', style='Normal')
    calc_book.add_paragraph(f'锚筋规格：{mj} 数量：{mjnums}', style='Normal')
    calc_book.add_paragraph(f'锚板规格：t{mbt}mm×{mbh}mm×{mbw}mm/Q345B', style='Normal')
    if vx > 0:  # 如果X向剪力大于0
        calc_book.add_paragraph('x向剪力设计值：', style='Normal')
        print('……', end=" ")
        mathtemp = r'V_x = ' + str(vx) + r'(N)'
        width = add_image(mathtemp, 'vx')
        calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/vx.jpg', width=Inches(width))
        calc_book.add_paragraph(f'沿x向剪力作用方向最外层锚筋中心线之间的距离：{zx}mm', style='Normal')
    if vz > 0:  # 如果z向剪力大于0
        calc_book.add_paragraph('z向剪力设计值：', style='Normal')
        print('……', end=" ")
        mathtemp = r'V_z = ' + str(vz) + r'(N)'
        width = add_image(mathtemp, 'vz')
        calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/vz.jpg', width=Inches(width))
        calc_book.add_paragraph(f'沿z向剪力作用方向最外层锚筋中心线之间的距离：{zz}mm', style='Normal')
    if vl > 0:
        calc_book.add_paragraph(f'剪力作用点距离锚板平面的距离：L={vl}mm', style='Normal')
    if mz > 0:
        calc_book.add_paragraph('x向剪力产生的弯矩：', style='Normal')
        mathtemp = r'M_z = V_x L =' + str(mz) + r'(N \cdot mm)'
        width = add_image(mathtemp, 'mz')
        calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/mz.jpg', width=Inches(width))
    if mx > 0:
        calc_book.add_paragraph('z向剪力产生的弯矩：', style='Normal')
        mathtemp = r'M_x = V_z L =' + str(mx) + r'(N \cdot mm)'
        width = add_image(mathtemp, 'mx')
        calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/mx.jpg', width=Inches(width))
    if ny > 0:  # ny大于0，是拉力
        calc_book.add_paragraph('法向拉力设计值：', style='Normal')
        print('……', end=" ")
        mathtemp = r'N = ' + str(ny) + r'(N)'
        width = add_image(mathtemp, 'ny')
        calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/ny.jpg', width=Inches(width))
    if ny < 0:  # ny小于0，是压力
        calc_book.add_paragraph('法向压力设计值：', style='Normal')
        print('……', end=" ")
        nyplus = -ny
        mathtemp = r'N = ' + str(nyplus) + r'(N)'
        width = add_image(mathtemp, 'ny')
        calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/ny.jpg', width=Inches(width))
        calc_book.add_paragraph('根据规范，法向压力设计值应满足下式', style='Normal')
        print('……', end=" ")
        fca = 0.5 * fc * mbh * mbw
        mathtemp = r'N \leq 0.5 f_c A =' + str(ceil(fca)) + r'(N)'
        width = add_image(mathtemp, 'fca')
        calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/fca.jpg', width=Inches(width))
        if nyplus <= fca:
            calc_book.add_paragraph('满足规范要求。', style='Quote').add_run('')
        else:
            calc_book.add_paragraph('', style='Quote').add_run('不满足规范要求。').font.color.rgb = RGBColor(0xff, 0x00, 0x00)

    calc_book.add_paragraph('锚筋的抗拉强度设计值：', style='Normal')
    print('……', end=" ")
    mathtemp = r'f_y = ' + str(fy) + r'(N/mm^2)'
    width = add_image(mathtemp, 'fy')
    calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/fy.jpg', width=Inches(width))

    calc_book.add_paragraph('混凝土轴心抗压强度设计值：', style='Normal')
    print('……', end=" ")
    mathtemp = r'f_c = ' + str(fc) + r'(N/mm^2)'
    width = add_image(mathtemp, 'fc')
    calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/fc.jpg', width=Inches(width))

    calc_book.add_paragraph('混凝土轴心抗拉强度设计值：', style='Normal')
    print('……', end=" ")
    mathtemp = r'f_t = ' + str(ft) + r'(N/mm^2)'
    width = add_image(mathtemp, 'ft')
    calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/ft.jpg', width=Inches(width))

    calc_book.add_paragraph('锚筋层数的影响系数：', style='Normal')
    print('……', end=" ")
    # 插入公式图片开始
    mathtemp = r'\alpha_r = ' + str(ar)
    width = add_image(mathtemp, 'ar')
    calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/ar.jpg', width=Inches(width))
    # 插入公式图片结束

    calc_book.add_paragraph('锚筋的受剪承载力系数：', style='Normal')
    print('……', end=" ")
    # 插入公式图片开始
    mathtemp = r'\alpha_v = (4.0 - 0.08d) \sqrt{\frac{f_c}{f_y}} = ' + str(round(av, 3))
    width = add_image(mathtemp, 'av')
    calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/av.jpg', width=Inches(width))
    # 插入公式图片结束

    calc_book.add_paragraph('锚板的弯曲变形折减系数：', style='Normal')
    print('……', end=" ")
    mathtemp = r'\alpha_b = 0.6 + 0.25 \frac{t}{d} = ' + str(round(ab, 3))
    width = add_image(mathtemp, 'ab')
    calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/ab.jpg', width=Inches(width))

    if ny >= 0:  # 受拉工况校核
        as1 = vx / (ar * av * fy) + vz / (ar * av * fy) + ny / (0.8 * ab * fy) + mx / (1.3 * ar * ab * fy * zz) + mz / (
                1.3 * ar * ab * fy * zx)
        as2 = ny / (0.8 * ab * fy) + mx / (0.4 * ar * ab * fy * zz) + mz / (0.4 * ar * ab * fy * zx)
        calc_book.add_paragraph('剪力、法向拉力和弯矩共同作用，锚筋总面积应满足下列两式', style='Normal')
        print('……', end=" ")
        mathtemp = r'A_s \geq \frac{V}{\alpha_r \alpha_v f_y} + \frac{N}{0.8 \alpha_b f_y} ' \
                   r'+ \frac{M}{1.3 \alpha_r \alpha_b f_y z}=' + str(ceil(as1)) + r'(mm^2)'
        width = add_image(mathtemp, 'as1')
        calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/as1.jpg', width=Inches(width))
        print('……', end=" ")
        mathtemp = r'A_s \geq \frac{N}{0.8 \alpha_b f_y} + \frac{M}{0.4 \alpha_r \alpha_b f_y z}=' + str(
            ceil(as2)) + r'(mm^2)'
        width = add_image(mathtemp, 'as2')
        calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/as2.jpg', width=Inches(width))
        calc_book.add_paragraph('锚筋实际总面积', style='Normal')
        print('……', end=" ")
        mathtemp = r'A_s =' + str(ceil(mjas)) + r'(mm^2)'
        width = add_image(mathtemp, 'as')
        calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/as.jpg', width=Inches(width))
        if mjas > as1 and mjas > as2:
            calc_book.add_paragraph('满足规范要求。', style='Quote')
        else:
            calc_book.add_paragraph('', style='Quote').add_run('不满足规范要求。').font.color.rgb = RGBColor(0xff, 0x00, 0x00)

    if ny < 0:  # 受压工况校核
        ny = -ny  # 临时将负数的压力转为正值，用于计算
        if 0 < mx < 0.4 * ny * zz:
            calc_book.add_paragraph('根据规范', style='Normal')
            mathtemp = r'M_x =' + str(mx) + r'(N \cdot mm) < 0.4 N z_z =' + str(ceil(0.4 * ny * zz)) + r'(N \cdot mm)'
            width = add_image(mathtemp, 'mnz1')
            calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/mnz1.jpg',
                                                                                    width=Inches(width))
            calc_book.add_paragraph('取', style='Normal')
            mathtemp = r'M_x = 0.4 N z_z =' + str(ceil(0.4 * ny * zz)) + r'(N \cdot mm)'
            width = add_image(mathtemp, 'mnz2')
            calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/mnz2.jpg',
                                                                                    width=Inches(width))
            mx = 0.4 * ny * zz
        if mx == 0:
            mx = 0.4 * ny * zz

        if 0 < mz < 0.4 * ny * zx:
            calc_book.add_paragraph('根据规范', style='Normal')
            mathtemp = r'M_z =' + str(mz) + r'(N \cdot mm) < 0.4 N z_x =' + str(ceil(0.4 * ny * zx)) + r'(N \cdot mm)'
            width = add_image(mathtemp, 'mnz3')
            calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/mnz3.jpg',
                                                                                    width=Inches(width))
            calc_book.add_paragraph('取', style='Normal')
            mathtemp = r'M_z = 0.4 N z_x =' + str(ceil(0.4 * ny * zx)) + r'(N \cdot mm)'
            width = add_image(mathtemp, 'mnz4')
            calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/mnz4.jpg',
                                                                                    width=Inches(width))
            mz = 0.4 * ny * zx
        if mz == 0:
            mz = 0.4 * ny * zx

        if vx == 0:
            vx = 0.3 * ny
        if vz == 0:
            vz = 0.3 * ny

        as3 = (vx - 0.3 * ny) / (ar * av * fy) + (vz - 0.3 * ny) / (ar * av * fy) + (mx - 0.4 * ny * zz) / (
                1.3 * ar * ab * fy * zz) + (mz - 0.4 * ny * zx) / (1.3 * ar * ab * fy * zx)
        as4 = (mx - 0.4 * ny * zz) / (0.4 * ar * ab * fy * zz) + (mz - 0.4 * ny * zx) / (0.4 * ar * ab * fy * zx)
        calc_book.add_paragraph('剪力、法向压力和弯矩共同作用，锚筋总面积应满足下列两式', style='Normal')
        mathtemp = r'A_s \geq \frac{V-0.3N}{\alpha_r \alpha_v f_y} + \frac{M-0.4Nz}{1.3 \alpha_r \alpha_b f_y z} =' + str(
            ceil(as3)) + r'(mm^2)'
        width = add_image(mathtemp, 'as3')
        calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/as3.jpg', width=Inches(width))
        print('……', end=" ")
        mathtemp = r'A_s \geq \frac{M-0.4Nz}{0.4 \alpha_r \alpha_b f_y z} =' + str(ceil(as4)) + r'(mm^2)'
        width = add_image(mathtemp, 'as4')
        calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/as4.jpg', width=Inches(width))
        calc_book.add_paragraph('锚筋实际总面积', style='Normal')
        print('……', end=" ")
        mathtemp = r'A_s =' + str(ceil(mjas)) + r'(mm^2)'
        width = add_image(mathtemp, 'as')
        calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/as.jpg', width=Inches(width))
        if mjas > as3 and mjas > as4:
            calc_book.add_paragraph('满足规范要求。', style='Quote')
        else:
            calc_book.add_paragraph('', style='Quote').add_run('不满足规范要求。').font.color.rgb = RGBColor(0xff, 0x00, 0x00)
        ny = -ny  # 恢复压力的负数值

    if ny >= 0:  # 受拉工况的锚固长度计算
        calc_book.add_paragraph('钢筋的外形系数：', style='Normal')
        print('……', end=" ")
        mathtemp = r'\alpha = 0.14'
        width = add_image(mathtemp, 'alpha')
        calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/alpha.jpg',
                                                                                width=Inches(width))
        calc_book.add_paragraph('基本锚固长度：', style='Normal')
        lab = 0.14 * fy * d / ft
        print('……', end=" ")
        mathtemp = r'l_{ab} = \alpha \frac{f_y}{f_t} d = ' + str(ceil(lab)) + '(mm)'
        width = add_image(mathtemp, 'lab')
        calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/lab.jpg', width=Inches(width))
        calc_book.add_paragraph('受拉钢筋的锚固长度限值：', style='Normal')
        mathtemp = r'l_{a} = \zeta_a l_{ab} = ' + str(ceil(lab * 1.1)) + '(mm)'
        width = add_image(mathtemp, 'la')
        calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/la.jpg', width=Inches(width))
        calc_book.add_paragraph('上式中，锚固长度修正系数按规范取：', style='Normal')
        mathtemp = r'\zeta_a = ' + str(1.1)
        width = add_image(mathtemp, 'zeta')
        calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/zeta.jpg', width=Inches(width))
        calc_book.add_paragraph('埋件处墙体厚度不满足锚固长度限值时采取双面锚板加强措施或采取搭接钢筋措施。', style='Normal')

    if ny < 0:  # 受压工况的锚固长度计算
        calc_book.add_paragraph('根据规范，受压直锚筋的锚固长度不应小于15d', style='Normal')
        calc_book.add_paragraph('受压钢筋的锚固长度限值：', style='Normal')
        mathtemp = r'l_{a} \geq 15d = ' + str(ceil(15 * d)) + '(mm)'
        width = add_image(mathtemp, 'lay')
        calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/lay.jpg', width=Inches(width))

    filename = f'{calc_i + 1}-埋件校核' + strftime("%Y-%m-%d-%H%M%S", localtime())
    calc_book.save(f'{filename}.docx')
    print(f'第{calc_i + 1}个计算书生成完成，保存在程序目录下，文件名为{filename}.docx')


print('计算书生成完成！')