from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import Mm
from docx.shared import Cm
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import urllib.parse
import urllib.request
from PIL import Image
from time import strftime
from time import localtime
from datetime import datetime
from docx.enum.text import WD_LINE_SPACING
from math import ceil
from os import mkdir
from os.path import exists
from math import pi

path = 'images'
if not exists(path):
    mkdir(path)


# 插入公式图片函数，参数1公式字符串，参数2公式图片的名称，返回值为公式图片的原始英寸宽度
def add_image(latex, pngname):
    math = urllib.parse.quote(latex)
    query_url = 'http://latex.xuming.science/latex-image.php?math=' + math
    try:
        chart = urllib.request.urlopen(query_url)
        f = open(f'{path}/{pngname}.png', 'wb')
        f.write(chart.read())
        f.close()
        img = Image.open(f'{path}/{pngname}.png')
        return img.size[0] / 96
    except:
        print('无法连接公式服务器，请检查网络连接或向软件作者提交问题')
        return 0


print('杆件稳定性校核自动生成word初始版')
print('')
print('20200525 by 徐明')
print('')

calc_book = Document()
# 设置正文字体
calc_book.styles['Normal'].font.name = 'Times New Roman'
calc_book.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
calc_book.styles['Normal'].font.size = Pt(12)
calc_book.styles['Normal'].font.color.rgb = RGBColor(0x00, 0x00, 0x00)
calc_book.styles['Normal'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
# 首行缩进，字符宽度等于字符高度，12pt=4.23mm, 1pt=0.3527mm
calc_book.styles['Normal'].paragraph_format.first_line_indent = Mm(8.46)
calc_book.styles['Normal'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
calc_book.styles['Normal'].paragraph_format.space_before = Mm(0)
calc_book.styles['Normal'].paragraph_format.space_after = Mm(0)
# 增加一个居中显示的样式
calc_book.styles['No Spacing'].font.name = 'Times New Roman'
calc_book.styles['No Spacing'].element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
calc_book.styles['No Spacing'].font.size = Pt(12)
calc_book.styles['No Spacing'].font.color.rgb = RGBColor(0x00, 0x00, 0x00)
calc_book.styles['No Spacing'].paragraph_format.first_line_indent = Mm(0)
calc_book.styles['No Spacing'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
calc_book.styles['No Spacing'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
calc_book.styles['No Spacing'].paragraph_format.space_before = Mm(0)
calc_book.styles['No Spacing'].paragraph_format.space_after = Mm(0)
# 设置标题1字体
calc_book.styles['Heading 1'].font.name = 'Times New Roman'
calc_book.styles['Heading 1'].element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
calc_book.styles['Heading 1'].element.rPr.rFonts.set(qn('w:asciiTheme'), 'Times New Roman')
calc_book.styles['Heading 1'].element.rPr.rFonts.set(qn('w:eastAsiaTheme'), '微软雅黑')
calc_book.styles['Heading 1'].element.rPr.rFonts.set(qn('w:hAnsiTheme'), 'Times New Roman')
calc_book.styles['Heading 1'].element.rPr.rFonts.set(qn('w:cstheme'), 'Times New Roman')
calc_book.styles['Heading 1'].font.size = Pt(16)
calc_book.styles['Heading 1'].font.color.rgb = RGBColor(0x00, 0x00, 0x00)
calc_book.styles['Heading 1'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
calc_book.styles['Heading 1'].paragraph_format.first_line_indent = Mm(0)
calc_book.styles['Heading 1'].paragraph_format.keep_with_next = True
calc_book.styles['Heading 1'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
calc_book.styles['Heading 1'].paragraph_format.space_before = Mm(0)
calc_book.styles['Heading 1'].paragraph_format.space_after = Mm(0)
# 设置标题2字体
calc_book.styles['Heading 2'].font.name = 'Times New Roman'
calc_book.styles['Heading 2'].element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
calc_book.styles['Heading 2'].element.rPr.rFonts.set(qn('w:asciiTheme'), 'Times New Roman')
calc_book.styles['Heading 2'].element.rPr.rFonts.set(qn('w:eastAsiaTheme'), '微软雅黑')
calc_book.styles['Heading 2'].element.rPr.rFonts.set(qn('w:hAnsiTheme'), 'Times New Roman')
calc_book.styles['Heading 2'].element.rPr.rFonts.set(qn('w:cstheme'), 'Times New Roman')
calc_book.styles['Heading 2'].font.size = Pt(14)
calc_book.styles['Heading 2'].font.color.rgb = RGBColor(0x00, 0x00, 0x00)
calc_book.styles['Heading 2'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
calc_book.styles['Heading 2'].paragraph_format.first_line_indent = Mm(0)
calc_book.styles['Heading 2'].paragraph_format.keep_with_next = True
calc_book.styles['Heading 2'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
calc_book.styles['Heading 2'].paragraph_format.space_before = Mm(0)
calc_book.styles['Heading 2'].paragraph_format.space_after = Mm(0)
# 设置标题3字体
calc_book.styles['Heading 3'].font.name = 'Times New Roman'
calc_book.styles['Heading 3'].element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
calc_book.styles['Heading 3'].element.rPr.rFonts.set(qn('w:asciiTheme'), 'Times New Roman')
calc_book.styles['Heading 3'].element.rPr.rFonts.set(qn('w:eastAsiaTheme'), '微软雅黑')
calc_book.styles['Heading 3'].element.rPr.rFonts.set(qn('w:hAnsiTheme'), 'Times New Roman')
calc_book.styles['Heading 3'].element.rPr.rFonts.set(qn('w:cstheme'), 'Times New Roman')
calc_book.styles['Heading 3'].font.size = Pt(14)
calc_book.styles['Heading 3'].font.color.rgb = RGBColor(0x00, 0x00, 0x00)
calc_book.styles['Heading 3'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
calc_book.styles['Heading 3'].paragraph_format.first_line_indent = Mm(0)
calc_book.styles['Heading 3'].paragraph_format.keep_with_next = True
calc_book.styles['Heading 3'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
calc_book.styles['Heading 3'].paragraph_format.space_before = Mm(0)
calc_book.styles['Heading 3'].paragraph_format.space_after = Mm(0)
# 增加一个右对齐的样式
calc_book.styles['Quote'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
calc_book.styles['Quote'].font.name = 'Italic'
calc_book.styles['Quote'].element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
calc_book.styles['Quote'].font.size = Pt(12)
calc_book.styles['Quote'].font.italic = False
calc_book.styles['Quote'].font.color.rgb = RGBColor(0x00, 0x00, 0x00)
calc_book.styles['Quote'].paragraph_format.space_before = Mm(0)
calc_book.styles['Quote'].paragraph_format.space_after = Mm(0)
calc_book.styles['Quote'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
calc_book.styles['Quote'].paragraph_format.first_line_indent = Mm(0)
'''
计算书文档属性设置
'''
print('文档属性初始化……')
# 文档标题
calc_book.core_properties.title = '计算书'
# 文档主题
calc_book.core_properties.subject = '计算书'
# 文档作者
calc_book.core_properties.author = 'xuming'
# 文档类别
calc_book.core_properties.category = 'calculation sheet'
# 文档注释
calc_book.core_properties.comments = 'Designed By Xuming. All Rights Reserved.'
# 文档创建时间
calc_book.core_properties.created = datetime.utcnow()
# 文档修改时间
calc_book.core_properties.modified = datetime.utcnow()

'''
计算书正文开始
'''

# 竖向撑杆参数
type1 = 2  # 截面属性，1为圆管，2为方管
d1 = 250
t1 = 10
l1 = 10400
n1 = 730000
xik1 = 0.825  # 钢号修正系数,Q235=1, Q345=0.825
fy1 = 345
# 许用参数
lam = 150
fp = 310

# b类截面轴向受压结构的稳定系数1-250范围
psi = [1, 1, 1, 0.999, 0.999, 0.998, 0.997, 0.996, 0.995, 0.994,
       0.992, 0.991, 0.989, 0.987, 0.985, 0.983, 0.981, 0.978, 0.976, 0.973,
       0.970, 0.967, 0.963, 0.960, 0.957, 0.953, 0.950, 0.946, 0.943, 0.939,
       0.936, 0.932, 0.929, 0.925, 0.921, 0.918, 0.914, 0.910, 0.906, 0.903,
       0.899, 0.895, 0.891, 0.886, 0.882, 0.878, 0.874, 0.870, 0.865, 0.861,
       0.856, 0.852, 0.847, 0.842, 0.837, 0.833, 0.828, 0.823, 0.818, 0.812,
       0.807, 0.802, 0.796, 0.791, 0.785, 0.780, 0.774, 0.768, 0.762, 0.757,
       0.751, 0.745, 0.738, 0.732, 0.726, 0.720, 0.713, 0.707, 0.701, 0.694,
       0.687, 0.681, 0.674, 0.668, 0.661, 0.654, 0.648, 0.641, 0.634, 0.628,
       0.621, 0.614, 0.607, 0.601, 0.594, 0.587, 0.581, 0.574, 0.568, 0.561,
       0.555, 0.548, 0.542, 0.535, 0.529, 0.523, 0.517, 0.511, 0.504, 0.498,
       0.492, 0.487, 0.481, 0.475, 0.469, 0.464, 0.458, 0.453, 0.447, 0.442,
       0.436, 0.431, 0.426, 0.421, 0.416, 0.411, 0.406, 0.401, 0.396, 0.392,
       0.387, 0.383, 0.378, 0.374, 0.369, 0.365, 0.361, 0.357, 0.352, 0.348,
       0.344, 0.340, 0.337, 0.333, 0.329, 0.325, 0.322, 0.318, 0.314, 0.311,
       0.308, 0.304, 0.301, 0.297, 0.294, 0.291, 0.288, 0.285, 0.282, 0.279,
       0.276, 0.273, 0.270, 0.267, 0.264, 0.262, 0.259, 0.256, 0.253, 0.251,
       0.248, 0.246, 0.243, 0.241, 0.238, 0.236, 0.234, 0.231, 0.229, 0.227,
       0.225, 0.222, 0.220, 0.218, 0.216, 0.214, 0.212, 0.210, 0.208, 0.206,
       0.204, 0.202, 0.200, 0.198, 0.196, 0.195, 0.193, 0.191, 0.189, 0.188,
       0.186, 0.184, 0.183, 0.181, 0.179, 0.178, 0.176, 0.175, 0.173, 0.172,
       0.170, 0.169, 0.167, 0.166, 0.164, 0.163, 0.162, 0.160, 0.159, 0.158,
       0.156, 0.155, 0.154, 0.152, 0.151, 0.150, 0.149, 0.147, 0.146, 0.145,
       0.144, 0.143, 0.142, 0.141, 0.139, 0.138, 0.137, 0.136, 0.135, 0.134,
       0.133, 0.132, 0.131, 0.130, 0.129, 0.128, 0.127, 0.126, 0.125, 0.124,
       0.123
       ]

# def psib(lam, fy):  # 公式计算压杆稳定系数
#     a1 = 0.65
#     a2 = 0.965
#     a3 = 0.3
#     es = 210000  # 钢材的弹性模量MPa
#     lamn = ((fy / es) ** 0.5) * lam / pi
#     if lamn > 0.215:
#         return ((a2 + a3 * lamn + (lamn ** 2)) - (((a2 + a3 * lamn + (lamn ** 2)) ** 2 - 4 * (lamn ** 2)) ** 0.5)) / (
#                 2 * (lamn ** 2))
#     else:
#         return 1 - a1 * (lamn ** 2)


calc_book.add_heading('五、杆件稳定性校核', level=1)
# calc_book.add_heading('5.1.竖向撑杆校核', level=2)
if type1 == 1:  # 截面属性，1为圆管，2为方管
    calc_book.add_paragraph(f'杆件截面为圆管：直径{d1}mm, 壁厚{t1}mm, 材质：Q345B', style='Normal')
    calc_book.add_paragraph('截面属性：', style='Normal')
    gxj = pi * ((d1 ** 4) - ((d1 - 2 * t1) ** 4)) / 64
    mj = pi * ((d1 ** 2) - ((d1 - 2 * t1) ** 2)) / 4
    hzbj = (gxj / mj) ** 0.5
    mass = mj * 1000 * 7850 * (10 ** (-9))
    lambda1 = l1 / hzbj
    sigma1 = n1 / mj
    mathtemp = r'I = ' + str(round(gxj, 0)) + r'(mm^4)'
    width = add_image(mathtemp, 'gxj')
    calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/gxj.png', width=Inches(width))
    mathtemp = r'i_x = ' + str(round(hzbj, 1)) + r'(mm)'
    width = add_image(mathtemp, 'hzbj')
    calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/hzbj.png', width=Inches(width))
    mathtemp = r'A = ' + str(round(mj, 1)) + r'(mm^2)'
    width = add_image(mathtemp, 'mj')
    calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/mj.png', width=Inches(width))
    calc_book.add_paragraph(f'计算长度：L={l1}(mm)', style='Normal')
    calc_book.add_paragraph(f'由前章节计算结果得最大轴向压力为：N={n1}(N)', style='Normal')
    calc_book.add_paragraph('长细比：', style='Normal')

    if lambda1 <= 150:
        mathtemp = r'\lambda = \frac{L}{i_x} = ' + str(ceil(lambda1)) + r'\leq [\lambda] =' + f'{lam}'
        width = add_image(mathtemp, 'lambda1')
        calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/lambda1.png',
                                                                                width=Inches(width))
        calc_book.add_paragraph('满足规范。', style='Quote')
    else:
        mathtemp = r'\lambda = \frac{L}{i_x} = ' + str(ceil(lambda1)) + r'> [\lambda]  =' + f'{lam}'
        width = add_image(mathtemp, 'lambda1')
        calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/lambda1.png',
                                                                                width=Inches(width))
        calc_book.add_paragraph('', style='Quote').add_run('不满足规范要求。').font.color.rgb = RGBColor(0xff, 0x00, 0x00)

    calc_book.add_paragraph('查表得稳定系数：', style='Normal')
    mathtemp = r'\psi = ' + str(psi[ceil(lambda1 / xik1)])
    width = add_image(mathtemp, 'psi')
    calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/psi.png',
                                                                            width=Inches(width))
    # calc_book.add_paragraph('用公式得稳定系数：', style='Normal')
    # mathtemp = r'\psi = ' + str(psib(lambda1, fy1))
    # width = add_image(mathtemp, 'psi1')
    # calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/psi1.png',
    #                                                                         width=Inches(width))
    calc_book.add_paragraph('轴向应力：', style='Normal')
    mathtemp = r'\sigma = \frac{N}{A} =' + str(round(n1 / mj, 1)) + r'(MPa)'
    width = add_image(mathtemp, 'sigma')
    calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/sigma.png',
                                                                            width=Inches(width))
    calc_book.add_paragraph('稳定性：', style='Normal')
    sigma2 = n1 / (psi[ceil(lambda1 / xik1)] * mj)
    if sigma2 < fp:
        mathtemp = r'\frac{N}{\psi A} =' + str(round(sigma2, 1)) + r'(MPa) < f =' + f'{fp}(MPa)'
        width = add_image(mathtemp, 'sigma2')
        calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/sigma2.png',
                                                                                width=Inches(width))
        calc_book.add_paragraph('满足规范。', style='Quote')
    else:
        mathtemp = r'\frac{N}{\psi A} =' + str(round(sigma2, 1)) + r'(MPa) > f =' + f'{fp}(MPa)'
        width = add_image(mathtemp, 'sigma2')
        calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/sigma2.png',
                                                                                width=Inches(width))
        calc_book.add_paragraph('', style='Quote').add_run('不满足规范要求。').font.color.rgb = RGBColor(0xff, 0x00, 0x00)
else:  # 方管截面校核
    calc_book.add_paragraph(f'杆件截面为方管：边长{d1}mm, 壁厚{t1}mm, 材质：Q345B', style='Normal')
    calc_book.add_paragraph('截面属性：', style='Normal')
    gxj = ((d1 ** 4) - ((d1 - 2 * t1) ** 4)) / 12
    mj = (d1 ** 2) - ((d1 - 2 * t1) ** 2)
    hzbj = (gxj / mj) ** 0.5
    mass = mj * 1000 * 7850 * (10 ** (-9))
    lambda1 = l1 / hzbj
    sigma1 = n1 / mj
    mathtemp = r'I = ' + str(round(gxj, 0)) + r'(mm^4)'
    width = add_image(mathtemp, 'gxj')
    calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/gxj.png', width=Inches(width))
    mathtemp = r'i_x = ' + str(round(hzbj, 1)) + r'(mm)'
    width = add_image(mathtemp, 'hzbj')
    calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/hzbj.png', width=Inches(width))
    mathtemp = r'A = ' + str(round(mj, 1)) + r'(mm^2)'
    width = add_image(mathtemp, 'mj')
    calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/mj.png', width=Inches(width))
    calc_book.add_paragraph(f'计算长度：L={l1}(mm)', style='Normal')
    calc_book.add_paragraph(f'由前章节计算结果得最大轴向压力为：N={n1}(N)', style='Normal')
    calc_book.add_paragraph('长细比：', style='Normal')

    if lambda1 <= 150:
        mathtemp = r'\lambda = \frac{L}{i_x} = ' + str(ceil(lambda1)) + r'\leq [\lambda] =' + f'{lam}'
        width = add_image(mathtemp, 'lambda1')
        calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/lambda1.png',
                                                                                width=Inches(width))
        calc_book.add_paragraph('满足规范。', style='Quote')
    else:
        mathtemp = r'\lambda = \frac{L}{i_x} = ' + str(ceil(lambda1)) + r'> [\lambda]  =' + f'{lam}'
        width = add_image(mathtemp, 'lambda1')
        calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/lambda1.png',
                                                                                width=Inches(width))
        calc_book.add_paragraph('', style='Quote').add_run('不满足规范要求。').font.color.rgb = RGBColor(0xff, 0x00, 0x00)

    calc_book.add_paragraph('查表得稳定系数：', style='Normal')
    mathtemp = r'\psi = ' + str(psi[ceil(lambda1 / xik1)])
    width = add_image(mathtemp, 'psi')
    calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/psi.png',
                                                                            width=Inches(width))
    # calc_book.add_paragraph('用公式得稳定系数：', style='Normal')
    # mathtemp = r'\psi = ' + str(psib(lambda1, fy1))
    # width = add_image(mathtemp, 'psi1')
    # calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/psi1.png',
    #                                                                         width=Inches(width))
    calc_book.add_paragraph('轴向应力：', style='Normal')
    mathtemp = r'\sigma = \frac{N}{A} =' + str(round(n1 / mj, 1)) + r'(MPa)'
    width = add_image(mathtemp, 'sigma')
    calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/sigma.png',
                                                                            width=Inches(width))
    calc_book.add_paragraph('稳定性：', style='Normal')
    sigma2 = n1 / (psi[ceil(lambda1 / xik1)] * mj)
    if sigma2 < fp:
        mathtemp = r'\frac{N}{\psi A} =' + str(round(sigma2, 1)) + r'(MPa) < f =' + f'{fp}(MPa)'
        width = add_image(mathtemp, 'sigma2')
        calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/sigma2.png',
                                                                                width=Inches(width))
        calc_book.add_paragraph('满足规范。', style='Quote')
    else:
        mathtemp = r'\frac{N}{\psi A} =' + str(round(sigma2, 1)) + r'(MPa) > f =' + f'{fp}(MPa)'
        width = add_image(mathtemp, 'sigma2')
        calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/sigma2.png',
                                                                                width=Inches(width))
        calc_book.add_paragraph('', style='Quote').add_run('不满足规范要求。').font.color.rgb = RGBColor(0xff, 0x00, 0x00)

'''
计算书结束，输出docx文档
'''
filename = f'杆件稳定性校核' + strftime("%Y-%m-%d-%H%M%S", localtime())
calc_book.save(f'{filename}.docx')
print(f'计算书生成结束，保存在程序目录下，文件名为{filename}.docx')
# input('按回车键退出......')
