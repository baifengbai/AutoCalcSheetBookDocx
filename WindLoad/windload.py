from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import Mm
from docx.shared import Cm
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import urllib.parse
import urllib.request
from PIL import Image
from time import strftime
from time import localtime
from datetime import datetime
from docx.enum.text import WD_LINE_SPACING
from math import ceil
from os import mkdir
from os.path import exists
from math import cos
from math import sin
from math import radians

print('======================================')
print('塔机抗风稳定校核计算书word生成器')
print('20200609 by 徐明')
print('适用工况：塔机随风自由旋转')
print('======================================')
print('计算书开始生成……')

path = 'images'
if not exists(path):
    mkdir(path)


# 插入公式图片函数，参数1公式字符串，参数2公式图片的名称，返回值为公式图片的原始英寸宽度
def add_image(latex, pngname):
    math = urllib.parse.quote(latex)
    query_url = 'http://latex.xuming.science/latex-image.php?math=' + math
    try:
        chart = urllib.request.urlopen(query_url)
        f = open(f'{path}/{pngname}.png', 'wb')
        f.write(chart.read())
        f.close()
        img = Image.open(f'{path}/{pngname}.png')
        return img.size[0] / 96
    except:
        print('无法连接公式服务器，请检查网络连接或向软件作者提交问题')
        return 0


# 插入公式图片函数，参数1公式字符串，参数2公式图片的名称，返回值为公式图片的高度mm
def add_image2(latex, pngname):
    math = urllib.parse.quote(latex)
    query_url = 'http://latex.xuming.science/latex-image.php?math=' + math
    try:
        chart = urllib.request.urlopen(query_url)
        f = open(f'{path}/{pngname}.png', 'wb')
        f.write(chart.read())
        f.close()
        img = Image.open(f'{path}/{pngname}.png')
        return img.size[1] * 25.4 / 96
    except:
        print('无法连接公式服务器，请检查网络连接或向软件作者提交问题')
        return 0


print('文档格式初始化……')
# 创建文档对象
calc_book = Document()
# 设置正文字体
calc_book.styles['Normal'].font.name = 'Times New Roman'
calc_book.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
calc_book.styles['Normal'].font.size = Pt(12)
calc_book.styles['Normal'].font.color.rgb = RGBColor(0x00, 0x00, 0x00)
calc_book.styles['Normal'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
# 首行缩进，字符宽度等于字符高度，12pt=4.23mm, 1pt=0.3527mm
font_height = 4.23
calc_book.styles['Normal'].paragraph_format.first_line_indent = Mm(0)
calc_book.styles['Normal'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
calc_book.styles['Normal'].paragraph_format.space_before = Mm(0)
calc_book.styles['Normal'].paragraph_format.space_after = Mm(0)

# 增加一个居中显示的样式
calc_book.styles['No Spacing'].paragraph_format.first_line_indent = Mm(0)
calc_book.styles['No Spacing'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
calc_book.styles['No Spacing'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
calc_book.styles['No Spacing'].paragraph_format.space_before = Mm(0)
calc_book.styles['No Spacing'].paragraph_format.space_after = Mm(0)
calc_book.styles['No Spacing'].font.size = Pt(10)
calc_book.styles['No Spacing'].font.name = 'Times New Roman'
calc_book.styles['No Spacing'].element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
# 增加一个左对齐的样式
calc_book.styles['Quote'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
calc_book.styles['Quote'].font.name = 'Times New Roman'
calc_book.styles['Quote'].element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
calc_book.styles['Quote'].font.size = Pt(10)
calc_book.styles['Quote'].font.italic = False
calc_book.styles['Quote'].font.color.rgb = RGBColor(0x00, 0x00, 0x00)
calc_book.styles['Quote'].paragraph_format.space_before = Mm(0)
calc_book.styles['Quote'].paragraph_format.space_after = Mm(0)
calc_book.styles['Quote'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
calc_book.styles['Quote'].paragraph_format.first_line_indent = Mm(0)
# 设置标题字体
calc_book.styles['Heading 1'].font.name = 'Times New Roman'
calc_book.styles['Heading 1'].element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
calc_book.styles['Heading 1'].element.rPr.rFonts.set(qn('w:asciiTheme'), 'Times New Roman')
calc_book.styles['Heading 1'].element.rPr.rFonts.set(qn('w:eastAsiaTheme'), '微软雅黑')
calc_book.styles['Heading 1'].element.rPr.rFonts.set(qn('w:hAnsiTheme'), 'Times New Roman')
calc_book.styles['Heading 1'].element.rPr.rFonts.set(qn('w:cstheme'), 'Times New Roman')
calc_book.styles['Heading 1'].font.size = Pt(22)
calc_book.styles['Heading 1'].font.color.rgb = RGBColor(0x00, 0x00, 0x00)
calc_book.styles['Heading 1'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
calc_book.styles['Heading 1'].paragraph_format.first_line_indent = Mm(0)
calc_book.styles['Heading 1'].paragraph_format.keep_with_next = True
calc_book.styles['Heading 1'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
calc_book.styles['Heading 1'].paragraph_format.space_before = Mm(0)
calc_book.styles['Heading 1'].paragraph_format.space_after = Mm(0)
# 设置标题2字体
calc_book.styles['Heading 2'].font.name = 'Times New Roman'
calc_book.styles['Heading 2'].element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
calc_book.styles['Heading 2'].element.rPr.rFonts.set(qn('w:asciiTheme'), 'Times New Roman')
calc_book.styles['Heading 2'].element.rPr.rFonts.set(qn('w:eastAsiaTheme'), '微软雅黑')
calc_book.styles['Heading 2'].element.rPr.rFonts.set(qn('w:hAnsiTheme'), 'Times New Roman')
calc_book.styles['Heading 2'].element.rPr.rFonts.set(qn('w:cstheme'), 'Times New Roman')
calc_book.styles['Heading 2'].font.size = Pt(18)
calc_book.styles['Heading 2'].font.color.rgb = RGBColor(0x00, 0x00, 0x00)
calc_book.styles['Heading 2'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
calc_book.styles['Heading 2'].paragraph_format.first_line_indent = Mm(0)
calc_book.styles['Heading 2'].paragraph_format.keep_with_next = True
calc_book.styles['Heading 2'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
calc_book.styles['Heading 2'].paragraph_format.space_before = Mm(0)
calc_book.styles['Heading 2'].paragraph_format.space_after = Mm(0)
# 设置标题3字体
calc_book.styles['Heading 3'].font.name = 'Times New Roman'
calc_book.styles['Heading 3'].element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
calc_book.styles['Heading 3'].element.rPr.rFonts.set(qn('w:asciiTheme'), 'Times New Roman')
calc_book.styles['Heading 3'].element.rPr.rFonts.set(qn('w:eastAsiaTheme'), '微软雅黑')
calc_book.styles['Heading 3'].element.rPr.rFonts.set(qn('w:hAnsiTheme'), 'Times New Roman')
calc_book.styles['Heading 3'].element.rPr.rFonts.set(qn('w:cstheme'), 'Times New Roman')
calc_book.styles['Heading 3'].font.size = Pt(14)
calc_book.styles['Heading 3'].font.color.rgb = RGBColor(0x00, 0x00, 0x00)
calc_book.styles['Heading 3'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
calc_book.styles['Heading 3'].paragraph_format.first_line_indent = Mm(0)
calc_book.styles['Heading 3'].paragraph_format.keep_with_next = True
calc_book.styles['Heading 3'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
calc_book.styles['Heading 3'].paragraph_format.space_before = Mm(0)
calc_book.styles['Heading 3'].paragraph_format.space_after = Mm(0)
'''
计算书文档属性设置
'''
print('文档属性初始化……')
# 文档标题
calc_book.core_properties.title = '计算书'
# 文档主题
calc_book.core_properties.subject = '计算书'
# 文档作者
calc_book.core_properties.author = 'xuming'
# 文档类别
calc_book.core_properties.category = 'calculation sheet'
# 文档注释
calc_book.core_properties.comments = 'Designed By Xuming. All Rights Reserved.'
# 文档创建时间
calc_book.core_properties.created = datetime.utcnow()
# 文档修改时间
calc_book.core_properties.modified = datetime.utcnow()

'''
计算书正文开始
'''

'''
参数定义
'''
# 计算书名称
jobname = '大疆ZSL1150抗风计算书'
# 塔机型号
tower_model = 'ZSL1150'
# 非工作10m高处基本风压
# pn = 1000
# 塔机最高处的计算高度
# ht = 220
# 风压高度变化系数, 保留两位小数
# kh = round(((((ht / 10) ** 0.14) + 0.4) / 1.4) ** 2, 2)

# 抗台风计算风压指定值
# 深圳防台风技术规程5.1.3中最大风压1870Pa
pnh = 1870
# 吊臂长度
beam_len = 58.5
# 非工作吊臂仰角
beam_ang = 60
# 塔机总力矩
tower_m = 1488
# 塔机后倾平衡力矩
tower_back = - ceil(tower_m * 0.45)
# 吊臂0度力矩
beam_0 = 702
# 吊臂非工作状态前倾力矩
beam_m = ceil(beam_0 * cos(radians(beam_ang)))
# 钩头重量
hook_mass = 1.3
# 钩头产生的前倾力矩
hook_m = ceil(hook_mass * beam_len * cos(radians(beam_ang)))
# 非工作状态下，上部结构不平衡
m0 = tower_back + beam_m + hook_m

# 吊臂
beam_d1 = 133  # 主弦直径mm
beam_d2 = 70  # 腹杆直径mm
beam_b = 2.25  # 主弦中心距m

beam_fg_len = 1.27 * beam_len * 2  # 腹杆长度总和
beam_a = beam_len * 2 * beam_d1 * 0.001 + beam_fg_len * beam_d2 * 0.001  # 特征面积
beam_phi = beam_a / (beam_len * (beam_b + beam_d1 * 0.001) * sin(radians(beam_ang)))  # 充实率
wind_v = (pnh / 0.625) ** 0.5  # 通过风压反推风速
re = 0.667 * wind_v * beam_d1 * 0.001  # 单位10^5
beam_lambda = round(beam_len * sin(radians(beam_ang)) / (beam_b + beam_d1 * 0.001), 1)

'''
文档生成
'''
calc_book.add_heading('2.1 无风状态塔机上部结构不平衡力矩', level=2)
calc_book.add_paragraph(f'吊臂长度：{beam_len}m', style='Normal')
calc_book.add_paragraph(f'非工作状态吊臂仰角：{beam_ang}°', style='Normal')
calc_book.add_paragraph(f'塔机后倾平衡力矩：{tower_back}t.m', style='Normal')
calc_book.add_paragraph(f'吊臂{beam_ang}°产生的前倾力矩：{beam_m}t.m', style='Normal')
calc_book.add_paragraph(f'钩头产生的前倾力矩：{hook_m}t.m', style='Normal')
calc_book.add_paragraph(f'非工作状态下塔机上部结构总不平衡力矩：M={tower_back}t.m+{beam_m}t.m+{hook_m}t.m={m0}t.m', style='Normal')
calc_book.add_paragraph('*为便于计算，规定由回转中心朝吊臂方向弯矩为正，由回转中心朝配重块方向弯矩为负。', style='Quote')

calc_book.add_heading('2.2 非工作状态风载荷', level=2)
mathtemp = r'F_{WN} = p_n(h) \times C \times A'
width = add_image(mathtemp, 'fwn')
calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/fwn.png', width=Inches(width))
calc_book.add_paragraph('式中：', style='Normal')

para1 = calc_book.add_paragraph('', style='Normal')
mathtemp = r'F_{WN}'
height = add_image2(mathtemp, 'fwn1')
if height > font_height:
    imgheight = font_height
else:
    imgheight = height
para1.add_run('').add_picture(f'{path}/fwn1.png', height=Mm(imgheight))  # Inches(width))
para1.add_run('——非工作状态垂直作用在所指构件纵轴线上的风载荷，单位为牛顿')
mathtemp = r'(N)'
height = add_image2(mathtemp, 'N')
if height > font_height:
    imgheight = font_height
else:
    imgheight = height
para1.add_run('').add_picture(f'{path}/N.png', height=Mm(imgheight))  # Inches(width))

para1 = calc_book.add_paragraph('', style='Normal')
mathtemp = r'p_n(h)'
height = add_image2(mathtemp, 'pnh')
if height > font_height:
    imgheight = font_height
else:
    imgheight = height
para1.add_run('').add_picture(f'{path}/pnh.png', height=Mm(imgheight))  # Inches(width))
para1.add_run('——高度h处的非工作状态计算风压，单位为牛顿每平方米')
mathtemp = r'(N/m^2)'
height = add_image2(mathtemp, 'Nm2')
if height > font_height:
    imgheight = font_height
else:
    imgheight = height
para1.add_run('').add_picture(f'{path}/Nm2.png', height=Mm(imgheight))  # Inches(width))

para1 = calc_book.add_paragraph('', style='Normal')
mathtemp = r'C'
height = add_image2(mathtemp, 'C')
if height > font_height:
    imgheight = font_height
else:
    imgheight = height
para1.add_run('').add_picture(f'{path}/C.png', height=Mm(imgheight))
para1.add_run('——所指构件的空气动力系数，与构件的特征面积A一起使用')

para1 = calc_book.add_paragraph('', style='Normal')
mathtemp = r'A'
height = add_image2(mathtemp, 'A')
if height > font_height:
    imgheight = font_height
else:
    imgheight = height
para1.add_run('').add_picture(f'{path}/A.png', height=Mm(imgheight))
para1.add_run('——所指构件的特征面积，单位为平方米')
mathtemp = r'(m^2)'
height = add_image2(mathtemp, 'm2')
if height > font_height:
    imgheight = font_height
else:
    imgheight = height
para1.add_run('').add_picture(f'{path}/m2.png', height=Mm(imgheight))  # Inches(width))

calc_book.add_heading('2.3 非工作状态计算风压', level=2)
calc_book.add_paragraph(f'{tower_model}塔机非工作状态计算风压取深圳防台风技术规程5.1.3中最大风压1870Pa。', style='Normal')
calc_book.add_paragraph('因此，非工作状态计算风压：', style='Normal')
mathtemp = r'p_n(h) = 1870 Pa'
width = add_image(mathtemp, 'pnh1870')
calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/pnh1870.png', width=Inches(width))

calc_book.add_heading('2.4 吊臂风载荷计算', level=2)
calc_book.add_paragraph(f'吊臂为空间桁架结构，主弦采用直径{beam_d1}mm圆管，腹杆采用直径{beam_d2}mm圆管，主弦中心距为{beam_b}m，吊臂总长{beam_len}m。',
                        style='Normal')
calc_book.add_paragraph('根据《GB/T 13752-2017 塔式起重机设计规范》 表B.5序号3及图B.8b，可得：', style='Normal')
para1 = calc_book.add_paragraph('特征面积', style='Normal')
mathtemp = r'A = ' + str(ceil(beam_a)) + 'm^2'
height = add_image2(mathtemp, 'beama')
if height > font_height:
    imgheight = font_height
else:
    imgheight = height
para1.add_run('').add_picture(f'{path}/beama.png', height=Mm(imgheight))

para1 = calc_book.add_paragraph('充实率', style='Normal')
mathtemp = r'\varphi = ' + str(round(beam_phi, 2))
height = add_image2(mathtemp, 'beamphi')
if height > font_height:
    imgheight = font_height
else:
    imgheight = height
para1.add_run('').add_picture(f'{path}/beamphi.png', height=Mm(imgheight))
print(f'充实率为{round(beam_phi, 2)}')

para1 = calc_book.add_paragraph('雷诺数', style='Normal')
mathtemp = r'Re = ' + str(round(re, 2)) + r'\times 10^5'
height = add_image2(mathtemp, 're')
if height > font_height:
    imgheight = font_height
else:
    imgheight = height
para1.add_run('').add_picture(f'{path}/re.png', height=Mm(imgheight))
print(f'雷诺数为{round(re, 2)}x10^5')

while True:
    try:
        beam_c0 = float(input("根据图B8b,得吊臂的空气动力系数C0: "))
        break
    except ValueError:
        print("输入错误，请输入正确数据")

para1 = calc_book.add_paragraph('空气动力系数', style='Normal')
mathtemp = r'C_0 = ' + str(round(beam_c0, 2))
height = add_image2(mathtemp, 'beam_c0')
if height > font_height:
    imgheight = font_height
else:
    imgheight = height
para1.add_run('').add_picture(f'{path}/beam_c0.png', height=Mm(imgheight))

calc_book.add_paragraph('表B.5  平面和空间格构式构件的特征面积和空气动力系数', style='No Spacing')
calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture('lib/B5.png', height=Cm(11.5))
calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture('lib/B8.png', height=Cm(4.6))
calc_book.add_paragraph('图B.8b 单根构件为圆形的空间格构式构件，其空气动力系数与雷诺数和充实率的关系', style='No Spacing')

calc_book.add_paragraph('吊臂的空气动力长细比：', style='Normal')
mathtemp = r'\lambda = \frac{l_a}{d} = \frac{' + str(round(beam_len * sin(radians(beam_ang)), 1)) + '}{' + str(
    beam_b + beam_d1 * 0.001) + '} = ' + str(beam_lambda)
width = add_image(mathtemp, 'beam_lambda')
calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/beam_lambda.png', width=Inches(width))

calc_book.add_paragraph('式中：', style='Normal')

para1 = calc_book.add_paragraph('', style='Normal')
mathtemp = r'\lambda'
height = add_image2(mathtemp, 'lambda')
if height > font_height:
    imgheight = font_height
else:
    imgheight = height
para1.add_run('').add_picture(f'{path}/lambda.png', height=Mm(imgheight))
para1.add_run('——空气动力长细比')

para1 = calc_book.add_paragraph('', style='Normal')
mathtemp = 'd'
height = add_image2(mathtemp, 'd')
if height > font_height:
    imgheight = font_height
else:
    imgheight = height
para1.add_run('').add_picture(f'{path}/d.png', height=Mm(imgheight))
para1.add_run('——构件的特征尺寸')

para1 = calc_book.add_paragraph('', style='Normal')
mathtemp = 'l_a'
height = add_image2(mathtemp, 'la')
if height > font_height:
    imgheight = font_height
else:
    imgheight = height
para1.add_run('').add_picture(f'{path}/la.png', height=Mm(imgheight))
para1.add_run('——构件的空气动力长度，按下式计算')

mathtemp = r'l_a = \alpha \times l = 1 \times' + str(beam_len) + r'\times sin(' + str(beam_ang) + r'^{\circ}) =' + str(
    round(beam_len * sin(radians(beam_ang)), 1))
width = add_image(mathtemp, 'la2')
calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/la2.png', width=Inches(width))

calc_book.add_paragraph('式中：', style='Normal')

para1 = calc_book.add_paragraph('', style='Normal')
mathtemp = r'l'
height = add_image2(mathtemp, 'l')
if height > font_height:
    imgheight = font_height
else:
    imgheight = height
para1.add_run('').add_picture(f'{path}/l.png', height=Mm(imgheight))
para1.add_run('——构件长度，即其两节点之间的距离')

para1 = calc_book.add_paragraph('', style='Normal')
mathtemp = r'\alpha'
# 行内图片公式，使用图片高度和字体高度的较小值
height = add_image2(mathtemp, 'alpha')
if height > font_height:
    imgheight = font_height
else:
    imgheight = height
para1.add_run('').add_picture(f'{path}/alpha.png', height=Mm(imgheight))
para1.add_run('——相对空气动力长度，按《GB/T 13752-2017 塔式起重机设计规范》取1')

print(f'充实率为{round(beam_phi, 2)}')
print(f'空气动力长细比为{beam_lambda}')
while True:
    try:
        beam_psi = float(input("根据图B1,得折减系数: "))
        break
    except ValueError:
        print("输入错误，请输入正确数据")

calc_book.add_paragraph('吊臂的空气动力系数：', style='Normal')
mathtemp = r'C = C_0 \times \Psi = ' + str(beam_c0) + r'\times' + str(beam_psi) + r'=' + str(
    round(beam_c0 * beam_psi, 3))
width = add_image(mathtemp, 'beam_c')
calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/beam_c.png', width=Inches(width))
calc_book.add_paragraph('式中：', style='Normal')

para1 = calc_book.add_paragraph('', style='Normal')
mathtemp = r'C_0'
# 行内图片公式，使用图片高度和字体高度的较小值
height = add_image2(mathtemp, 'c0')
if height > font_height:
    imgheight = font_height
else:
    imgheight = height
para1.add_run('').add_picture(f'{path}/c0.png', height=Mm(imgheight))
para1.add_run('——无限长直的等截面构件的空气动力系数')

para1 = calc_book.add_paragraph('', style='Normal')
mathtemp = r'\Psi'
# 行内图片公式，使用图片高度和字体高度的较小值
height = add_image2(mathtemp, 'psi')
if height > font_height:
    imgheight = font_height
else:
    imgheight = height
para1.add_run('').add_picture(f'{path}/psi.png', height=Mm(imgheight))
para1.add_run('——折减系数，为适合有限长构件而折减').add_picture(f'{path}/c0.png', height=Mm(imgheight))
para1.add_run('。').add_picture(f'{path}/psi.png', height=Mm(imgheight))
para1.add_run('与单根构件的空气动力长细比有关；如果是格构式构件，还与其充实率有关。按图B.1取值')
mathtemp = r'\Psi =' + str(beam_psi)
# 行内图片公式，使用图片高度和字体高度的较小值
height = add_image2(mathtemp, 'psi2')
if height > font_height:
    imgheight = font_height
else:
    imgheight = height
para1.add_run('').add_picture(f'{path}/psi2.png', height=Mm(imgheight))

calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture('lib/B1.png', height=Cm(7.2))
calc_book.add_paragraph('图B.1 折减系数与空气动力长细比、结构充实率之间的关系', style='No Spacing')



'''
计算书结束，输出docx文档
'''
filename = f'{jobname}' + strftime("%Y-%m-%d-%H%M%S", localtime())
calc_book.save(f'{filename}.docx')
print(f'计算书生成结束，保存在程序目录下，文件名为{filename}.docx')
