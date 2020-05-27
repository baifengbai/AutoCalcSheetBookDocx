from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import Mm
from docx.shared import Cm
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import urllib.parse
import urllib.request
from PIL import Image
from time import strftime
from time import localtime
from datetime import datetime
from docx.enum.text import WD_LINE_SPACING
from math import ceil
from os import mkdir
from os.path import exists
import pandas as pd
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT


calc_book = Document()

# 设置正文字体
calc_book.styles['Normal'].font.name = 'Times New Roman'
calc_book.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
calc_book.styles['Normal'].font.size = Pt(12)
calc_book.styles['Normal'].font.color.rgb = RGBColor(0x00, 0x00, 0x00)
calc_book.styles['Normal'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
# 首行缩进，字符宽度等于字符高度，12pt=4.23mm, 1pt=0.3527mm
calc_book.styles['Normal'].paragraph_format.first_line_indent = Mm(8.46)
calc_book.styles['Normal'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
calc_book.styles['Normal'].paragraph_format.space_before = Mm(0)
calc_book.styles['Normal'].paragraph_format.space_after = Mm(0)

calc_book.add_paragraph('8工况结果汇总', style='Normal')


rows_num = 64
cols_num = 4

mytable = calc_book.add_table(rows=rows_num, cols=cols_num, style='Table Grid')
mytable.alignment = WD_TABLE_ALIGNMENT.CENTER
mytable.autofit = False

skiplist = [0, 1, 2, 25, 26, 49, 50, 73, 74, 97, 98, 121, 122, 145, 146, 169, 170,
            193, 194, 195, 218, 219, 242, 243, 266]
ansys = pd.read_table('TheResult.txt', sep='|', skiprows=skiplist, header=None)

for r in range(8):
    mytable.cell(0 + r * 8, 0).text = f'工况-{r + 1}'
    mytable.cell(1 + r * 8, 0).text = '节点'
    mytable.cell(1 + r * 8, 1).text = 'F_X(t)'
    mytable.cell(1 + r * 8, 2).text = 'F_y(t)'
    mytable.cell(1 + r * 8, 3).text = 'F_z(t)'
    for i in range(6):
        mytable.cell(2 + i + r * 8, 0).text = str(int(ansys.iloc[i + r * 22, 0]))
        for j in range(1, 4):
            mytable.cell(2 + i + r * 8, j).text = str(ansys.iloc[i + r * 22, j])

# 设置单元格格式
for c in range(cols_num):
    for cell in mytable.columns[c].cells:
        cell.width = Cm(3.5)
        cell.paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        cell.paragraphs[0].paragraph_format.first_line_indent = Mm(0)
        cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
for row in mytable.rows:
    row.height = Cm(0.8)
# mytable.columns[0].cells[0].paragraphs[0].paragraph_format.space_after = Inches(0)
# mytable.columns[0].cells[0].paragraphs[0].paragraph_format.space_before = Inches(0)
# mytable.columns[0].cells[0].paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
# mytable.columns[0].cells[0].paragraphs[0].paragraph_format.first_line_indent = Mm(0)
# mytable.columns[0].cells[0].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# print(mytable.columns[0].cells[0].paragraphs[0].paragraph_format._element.xml)
calc_book.add_paragraph('', style='Normal')
calc_book.add_paragraph('杆件轴力包络值', style='Normal')
mytable2 = calc_book.add_table(rows=5, cols=3, style='Table Grid')
mytable2.alignment = WD_TABLE_ALIGNMENT.CENTER
mytable2.autofit = False
mytable2.cell(0, 0).text = '名称'
# mytable2.cell(0, 1).text = '杆件长度/mm'
mytable2.cell(0, 1).text = '轴向力/t'

mytable2.cell(1, 0).text = '水平杆1'

mytable2.cell(1, 1).text = str(ceil(ansys.iloc[6 + 10 * 22, 1]))

mytable2.cell(2, 0).text = '水平杆2'

mytable2.cell(2, 1).text = str(ceil(ansys.iloc[7 + 10 * 22, 1]))

mytable2.cell(3, 0).text = '竖向撑杆1'

mytable2.cell(3, 1).text = str(ceil(ansys.iloc[8 + 10 * 22, 1]))

mytable2.cell(4, 0).text = '竖向撑杆2'
mytable2.cell(4, 1).text = str(ceil(ansys.iloc[9 + 10 * 22, 1]))

# 设置单元格格式
for c in range(3):
    for cell in mytable2.columns[c].cells:
        cell.width = Cm(3.5)
        cell.paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        cell.paragraphs[0].paragraph_format.first_line_indent = Mm(0)
        cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
for row in mytable2.rows:
    row.height = Cm(0.8)

filename = f'test' + strftime("%Y-%m-%d-%H%M%S", localtime())
calc_book.save(f'{filename}.docx')
