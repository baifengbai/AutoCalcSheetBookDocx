from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import Mm
from docx.shared import Cm
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import urllib.parse
import urllib.request
from PIL import Image
from time import strftime
from time import localtime
from datetime import datetime
from docx.enum.text import WD_LINE_SPACING
from math import ceil
from os import mkdir
from os.path import exists


print('外挂支持8工况应力云图自动提取到word初始版')
print('')
print('20200525 by 徐明')
print('')

calc_book = Document()
# 设置正文字体
calc_book.styles['Normal'].font.name = 'Times New Roman'
calc_book.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
calc_book.styles['Normal'].font.size = Pt(12)
calc_book.styles['Normal'].font.color.rgb = RGBColor(0x00, 0x00, 0x00)
calc_book.styles['Normal'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
# 首行缩进，字符宽度等于字符高度，12pt=4.23mm, 1pt=0.3527mm
calc_book.styles['Normal'].paragraph_format.first_line_indent = Mm(8.46)
calc_book.styles['Normal'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
calc_book.styles['Normal'].paragraph_format.space_before = Mm(0)
calc_book.styles['Normal'].paragraph_format.space_after = Mm(0)
# 增加一个居中显示的样式
calc_book.styles['No Spacing'].font.name = 'Times New Roman'
calc_book.styles['No Spacing'].element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
calc_book.styles['No Spacing'].font.size = Pt(12)
calc_book.styles['No Spacing'].font.color.rgb = RGBColor(0x00, 0x00, 0x00)
calc_book.styles['No Spacing'].paragraph_format.first_line_indent = Mm(0)
calc_book.styles['No Spacing'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
calc_book.styles['No Spacing'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
calc_book.styles['No Spacing'].paragraph_format.space_before = Mm(0)
calc_book.styles['No Spacing'].paragraph_format.space_after = Mm(0)
# 设置标题1字体
calc_book.styles['Heading 1'].font.name = 'Times New Roman'
calc_book.styles['Heading 1'].element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
calc_book.styles['Heading 1'].element.rPr.rFonts.set(qn('w:asciiTheme'), 'Times New Roman')
calc_book.styles['Heading 1'].element.rPr.rFonts.set(qn('w:eastAsiaTheme'), '微软雅黑')
calc_book.styles['Heading 1'].element.rPr.rFonts.set(qn('w:hAnsiTheme'), 'Times New Roman')
calc_book.styles['Heading 1'].element.rPr.rFonts.set(qn('w:cstheme'), 'Times New Roman')
calc_book.styles['Heading 1'].font.size = Pt(16)
calc_book.styles['Heading 1'].font.color.rgb = RGBColor(0x00, 0x00, 0x00)
calc_book.styles['Heading 1'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
calc_book.styles['Heading 1'].paragraph_format.first_line_indent = Mm(0)
calc_book.styles['Heading 1'].paragraph_format.keep_with_next = True
calc_book.styles['Heading 1'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
calc_book.styles['Heading 1'].paragraph_format.space_before = Mm(0)
calc_book.styles['Heading 1'].paragraph_format.space_after = Mm(0)
# 设置标题2字体
calc_book.styles['Heading 2'].font.name = 'Times New Roman'
calc_book.styles['Heading 2'].element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
calc_book.styles['Heading 2'].element.rPr.rFonts.set(qn('w:asciiTheme'), 'Times New Roman')
calc_book.styles['Heading 2'].element.rPr.rFonts.set(qn('w:eastAsiaTheme'), '微软雅黑')
calc_book.styles['Heading 2'].element.rPr.rFonts.set(qn('w:hAnsiTheme'), 'Times New Roman')
calc_book.styles['Heading 2'].element.rPr.rFonts.set(qn('w:cstheme'), 'Times New Roman')
calc_book.styles['Heading 2'].font.size = Pt(14)
calc_book.styles['Heading 2'].font.color.rgb = RGBColor(0x00, 0x00, 0x00)
calc_book.styles['Heading 2'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
calc_book.styles['Heading 2'].paragraph_format.first_line_indent = Mm(0)
calc_book.styles['Heading 2'].paragraph_format.keep_with_next = True
calc_book.styles['Heading 2'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
calc_book.styles['Heading 2'].paragraph_format.space_before = Mm(0)
calc_book.styles['Heading 2'].paragraph_format.space_after = Mm(0)
# 设置标题3字体
calc_book.styles['Heading 3'].font.name = 'Times New Roman'
calc_book.styles['Heading 3'].element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
calc_book.styles['Heading 3'].element.rPr.rFonts.set(qn('w:asciiTheme'), 'Times New Roman')
calc_book.styles['Heading 3'].element.rPr.rFonts.set(qn('w:eastAsiaTheme'), '微软雅黑')
calc_book.styles['Heading 3'].element.rPr.rFonts.set(qn('w:hAnsiTheme'), 'Times New Roman')
calc_book.styles['Heading 3'].element.rPr.rFonts.set(qn('w:cstheme'), 'Times New Roman')
calc_book.styles['Heading 3'].font.size = Pt(14)
calc_book.styles['Heading 3'].font.color.rgb = RGBColor(0x00, 0x00, 0x00)
calc_book.styles['Heading 3'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
calc_book.styles['Heading 3'].paragraph_format.first_line_indent = Mm(0)
calc_book.styles['Heading 3'].paragraph_format.keep_with_next = True
calc_book.styles['Heading 3'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
calc_book.styles['Heading 3'].paragraph_format.space_before = Mm(0)
calc_book.styles['Heading 3'].paragraph_format.space_after = Mm(0)

'''
计算书文档属性设置
'''
print('文档属性初始化……')
# 文档标题
calc_book.core_properties.title = '计算书'
# 文档主题
calc_book.core_properties.subject = '计算书'
# 文档作者
calc_book.core_properties.author = 'xuming'
# 文档类别
calc_book.core_properties.category = 'calculation sheet'
# 文档注释
calc_book.core_properties.comments = 'Designed By Xuming. All Rights Reserved.'
# 文档创建时间
calc_book.core_properties.created = datetime.utcnow()
# 文档修改时间
calc_book.core_properties.modified = datetime.utcnow()

'''
计算书正文开始
'''
# calc_book.add_heading('四、外挂支撑系统有限元建模分析', level=1)
# calc_book.add_paragraph('采用大型通用有限元软件ANSYS按“三、外挂支撑系统图纸”建模，按“二、结构反力”加载分析', style='Normal')
# calc_book.add_paragraph('取塔机吊臂旋转每隔45度为一个工况，分析8种工况。', style='Normal')

titlist = '4.3'
picpath = '850-3'
picnum = 850000

calc_book.add_heading(f'{titlist}.ZSL850外挂支撑系统', level=2)
calc_book.add_paragraph('', style='Normal')
calc_book.add_paragraph('外挂架模型', style='No Spacing')
calc_book.add_paragraph('', style='Normal')
calc_book.add_paragraph('8种工况示意图', style='No Spacing')

for i in range(1,8):
    calc_book.add_heading(f'{titlist}.{i}.工况{i}', level=3)
    calc_book.add_paragraph('外挂架X向最大挠度00mm，Y向最大挠度00mm，Z向最大挠度00mm:', style='Normal')
    calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{picpath}/{picnum}.png', height=Cm(7))
    picnum = picnum + 1
    calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{picpath}/{picnum}.png', height=Cm(7))
    picnum = picnum + 1
    calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{picpath}/{picnum}.png', height=Cm(7))
    picnum = picnum + 1
    calc_book.add_paragraph('主梁最大Mises应力00MPa:', style='Normal')
    calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{picpath}/{picnum}.png', height=Cm(7))
    picnum = picnum + 1
    calc_book.add_paragraph('竖向撑杆最大Mises应力00MPa:', style='Normal')
    calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{picpath}/{picnum}.png', height=Cm(7))
    picnum = picnum + 1
    calc_book.add_paragraph('水平撑杆最大Mises应力00MPa:', style='Normal')
    calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{picpath}/{picnum}.png', height=Cm(7))
    picnum = picnum + 1

'''
计算书结束，输出docx文档
'''
filename = f'{titlist}外挂支撑系统' + strftime("%Y-%m-%d-%H%M%S", localtime())
calc_book.save(f'{filename}.docx')
print(f'计算书生成结束，保存在程序目录下，文件名为{filename}.docx')
input('按回车键退出......')





