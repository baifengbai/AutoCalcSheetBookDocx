from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import Mm
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import urllib.parse
import urllib.request
from PIL import Image
from time import strftime
from time import localtime
from datetime import datetime
from docx.enum.text import WD_LINE_SPACING
from math import ceil
from os import mkdir
from os.path import exists

print('======================================')
print('=========墙体埋件计算书word生成器=========')
print('=========内测版本20200522 by 徐明=======')
print('======================================')
print('计算书开始生成……')

path = 'images'
if not exists(path):
    mkdir(path)


# 插入公式图片函数，参数1公式字符串，参数2公式图片的名称，返回值为公式图片的原始英寸宽度
def add_image(latex, pngname):
    math = urllib.parse.quote(latex)
    query_url = 'http://latex.xuming.science/latex-image.php?math=' + math
    try:
        chart = urllib.request.urlopen(query_url)
        f = open(f'{path}/{pngname}.png', 'wb')
        f.write(chart.read())
        f.close()
        img = Image.open(f'{path}/{pngname}.png')
        return img.size[0] / 96
    except:
        print('无法连接公式服务器，请检查网络连接或向软件作者提交问题')
        return 0
    # chart = urllib.request.urlopen(query_url)
    # f = open(f"{pngname}.png", "wb")
    # f.write(chart.read())
    # f.close()
    # img = Image.open(f"{pngname}.png")
    # return img.size[0] / 96


print('文档格式初始化……')
# 创建文档对象
calc_book = Document()
# 设置正文字体
calc_book.styles['Normal'].font.name = 'Italic'
calc_book.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
calc_book.styles['Normal'].font.size = Pt(12)
calc_book.styles['Normal'].font.color.rgb = RGBColor(0x00, 0x00, 0x00)
calc_book.styles['Normal'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
# 首行缩进，字符宽度等于字符高度，12pt=4.23mm, 1pt=0.3527mm
calc_book.styles['Normal'].paragraph_format.first_line_indent = Mm(8.46)
calc_book.styles['Normal'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
calc_book.styles['Normal'].paragraph_format.space_before = Mm(0)
calc_book.styles['Normal'].paragraph_format.space_after = Mm(0)
# 增加一个居中显示的样式
calc_book.styles['No Spacing'].paragraph_format.first_line_indent = Mm(0)
calc_book.styles['No Spacing'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
calc_book.styles['No Spacing'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
calc_book.styles['No Spacing'].paragraph_format.space_before = Mm(0)
calc_book.styles['No Spacing'].paragraph_format.space_after = Mm(0)
# 增加一个右对齐的样式
calc_book.styles['Quote'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
calc_book.styles['Quote'].font.name = 'Italic'
calc_book.styles['Quote'].element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
calc_book.styles['Quote'].font.size = Pt(12)
calc_book.styles['Quote'].font.italic = False
calc_book.styles['Quote'].font.color.rgb = RGBColor(0x00, 0x00, 0x00)
calc_book.styles['Quote'].paragraph_format.space_before = Mm(0)
calc_book.styles['Quote'].paragraph_format.space_after = Mm(0)
calc_book.styles['Quote'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
calc_book.styles['Quote'].paragraph_format.first_line_indent = Mm(0)
# 设置标题字体
calc_book.styles['Heading 1'].font.name = 'Times New Roman'
calc_book.styles['Heading 1'].element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
calc_book.styles['Heading 1'].element.rPr.rFonts.set(qn('w:asciiTheme'), 'Times New Roman')
calc_book.styles['Heading 1'].element.rPr.rFonts.set(qn('w:eastAsiaTheme'), '微软雅黑')
calc_book.styles['Heading 1'].element.rPr.rFonts.set(qn('w:hAnsiTheme'), 'Times New Roman')
calc_book.styles['Heading 1'].element.rPr.rFonts.set(qn('w:cstheme'), 'Times New Roman')
calc_book.styles['Heading 1'].font.size = Pt(22)
calc_book.styles['Heading 1'].font.color.rgb = RGBColor(0x00, 0x00, 0x00)
calc_book.styles['Heading 1'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
calc_book.styles['Heading 1'].paragraph_format.first_line_indent = Mm(0)
calc_book.styles['Heading 1'].paragraph_format.keep_with_next = True
calc_book.styles['Heading 1'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
calc_book.styles['Heading 1'].paragraph_format.space_before = Mm(0)
calc_book.styles['Heading 1'].paragraph_format.space_after = Mm(0)

'''
计算书文档属性设置
'''
print('文档属性初始化……')
# 文档标题
calc_book.core_properties.title = '计算书'
# 文档主题
calc_book.core_properties.subject = '计算书'
# 文档作者
calc_book.core_properties.author = 'xuming'
# 文档类别
calc_book.core_properties.category = 'calculation sheet'
# 文档注释
calc_book.core_properties.comments = 'Designed By Xuming. All Rights Reserved.'
# 文档创建时间
calc_book.core_properties.created = datetime.utcnow()
# 文档修改时间
calc_book.core_properties.modified = datetime.utcnow()

'''
计算书正文开始
'''
print('请输入参数')

# =========常规参数输入====================================
while True:
    try:
        d = int(input("锚筋直径(mm): "))
        break
    except ValueError:
        print("输入错误，请输入正确的锚筋直径数字[整数]")

# d = 28  # 锚筋直径，单位mm
mj = f'HRB400-ø{d}'  # 锚筋规格
# while True:
#     try:
#         mjnums = int(input("锚筋数量: "))
#         break
#     except ValueError:
#         print("输入错误，请输入正确的锚筋数量（整数）")

# mjnums = 35  # 锚筋数量
while True:
    try:
        mbt = int(input("锚板厚度(mm): "))
        break
    except ValueError:
        print("输入错误，请输入正确的锚板厚度[整数]")
# mbt = 40  # 锚板厚度，单位mm
while True:
    try:
        mbh = int(input("锚板高度(mm)[竖向]: "))
        break
    except ValueError:
        print("输入错误，请输入正确的锚板高度[整数]")
# mbh = 1344  # 锚板高度，单位mm
while True:
    try:
        mbw = int(input("锚板宽度(mm)[水平向]: "))
        break
    except ValueError:
        print("输入错误，请输入正确的锚板宽度[整数]")
# mbw = 1008  # 锚板宽度，单位mm
while True:
    try:
        mjjj = int(input("锚筋间距(mm)[双向等间距]: "))
        break
    except ValueError:
        print("输入错误，请输入正确数据[整数]")
while True:
    try:
        xrow = int(input("水平锚筋排数: "))
        break
    except ValueError:
        print("输入错误，请输入正确数据[整数]")
while True:
    try:
        zrow = int(input("竖向锚筋排数: "))
        break
    except ValueError:
        print("输入错误，请输入正确数据[整数]")
while True:
    try:
        fy = int(input("锚筋的抗拉强度设计值(N/mm^2)[不大于300，一般计算取300]: "))
        break
    except ValueError:
        print("输入错误，请输入正确数据[整数]")
# fy = 300  # 锚筋的抗拉强度设计值，单位N/mm^2
while True:
    try:
        fc = float(input("混凝土轴心抗压强度设计值(N/mm^2)[C40,取19.1,其它情况查规范]: "))
        break
    except ValueError:
        print("输入错误，请输入正确数据")
# fc = 19.1  # 混凝土轴心抗压强度设计值，单位N/mm^2
while True:
    try:
        ft = float(input("混凝土轴心抗拉强度设计值(N/mm^2)[C40,取1.71,其它情况查规范]: "))
        break
    except ValueError:
        print("输入错误，请输入正确数据")
# ft = 1.71  # 混凝土轴心抗拉强度设计值，单位N/mm^2
while True:
    try:
        vx = int(input("x向剪力(N)[水平方向,填正值,没有则填0]: "))
        break
    except ValueError:
        print("输入错误，请输入正确数据[整数]")
# vx = 356000  # x向剪力，单位N，这个方向没有剪力就填0
while True:
    try:
        vz = int(input("z向剪力(N)[竖向,填正值,没有则填0]: "))
        break
    except ValueError:
        print("输入错误，请输入正确数据[整数]")
# vz = 507000  # z向剪力，单位N，这个方向没有剪力就填0
while True:
    try:
        ny = int(input("法向拉力（正）或法向压力（负）(N)[没有则填0]: "))
        break
    except ValueError:
        print("输入错误，请输入正确数据[整数]")
# ny = 1618000  # 法向拉力（正）或法向压力（负），单位N
zx = mjjj * (xrow - 1)  # x向剪力作用方向最外层锚筋中心线之间的距离，单位mm
zz = mjjj * (zrow - 1)  # z向剪力作用方向最外层锚筋中心线之间的距离，单位mm
row = max(xrow, zrow)  # 锚筋排数取大值
while True:
    try:
        vl = int(input("剪力作用点距离锚板平面的距离(mm): "))
        break
    except ValueError:
        print("输入错误，请输入正确数据[整数]")
# vl = 200  # 剪力作用点距离锚板平面的距离，单位mm
# =======================================================
mjnums = xrow * zrow  # 锚筋数量
mz = vx * vl  # z向弯矩设计值，由x向剪力产生，单位N.mm
mx = vz * vl  # x向弯矩设计值，由z向剪力产生，单位N.mm

if row == 2:
    ar = 1
elif row == 3:
    ar = 0.9
elif row == 4:
    ar = 0.85
else:
    ar = 0.85

# 锚筋层数的影响系数；当锚筋按等间距布置时：两层取1.0；三层取0.9；四层取0.85

av = (4 - 0.08 * d) * ((fc / fy) ** 0.5)  # 锚筋的受剪承载力系数
if av > 0.7:  # 根据规范大于0.7时取0.7
    av = 0.7
ab = 0.6 + 0.25 * mbt / d  # 锚板的弯曲变形折减系数
mjas = mjnums * 3.14 * ((d / 2) ** 2)  # 锚筋的总面积

# calc_book.add_heading('1. 埋件校核', level=1)
calc_book.add_paragraph('计算依据：GB50010-2010混凝土设计规范(2015年版)', style='Normal')
calc_book.add_paragraph(f'锚筋规格：{mj} 数量：{mjnums}', style='Normal')
calc_book.add_paragraph(f'锚板规格：t{mbt}mm×{mbh}mm×{mbw}mm/Q345B', style='Normal')
if vx > 0:
    calc_book.add_paragraph('x向剪力设计值：', style='Normal')
    print('……', end=" ")
    mathtemp = r'V_x = ' + str(vx) + r'(N)'
    width = add_image(mathtemp, 'vx')
    calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/vx.png', width=Inches(width))
    calc_book.add_paragraph(f'沿x向剪力作用方向最外层锚筋中心线之间的距离：{zx}mm', style='Normal')

if vz > 0:
    calc_book.add_paragraph('z向剪力设计值：', style='Normal')
    print('……', end=" ")
    mathtemp = r'V_z = ' + str(vz) + r'(N)'
    width = add_image(mathtemp, 'vz')
    calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/vz.png', width=Inches(width))
    calc_book.add_paragraph(f'沿z向剪力作用方向最外层锚筋中心线之间的距离：{zz}mm', style='Normal')

calc_book.add_paragraph(f'剪力作用点距离锚板平面的距离：L={vl}mm', style='Normal')
if vx > 0:
    calc_book.add_paragraph('x向剪力产生的弯矩：', style='Normal')
    mathtemp = r'M_z = V_x L =' + str(mz) + r'(N \cdot mm)'
    width = add_image(mathtemp, 'mz')
    calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/mz.png', width=Inches(width))
if vz > 0:
    calc_book.add_paragraph('z向剪力产生的弯矩：', style='Normal')
    mathtemp = r'M_x = V_z L =' + str(mx) + r'(N \cdot mm)'
    width = add_image(mathtemp, 'mx')
    calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/mx.png', width=Inches(width))

if ny >= 0:
    calc_book.add_paragraph('法向拉力设计值：', style='Normal')
    print('……', end=" ")
    mathtemp = r'N = ' + str(ny) + r'(N)'
    width = add_image(mathtemp, 'ny')
    calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/ny.png', width=Inches(width))

if ny < 0:
    calc_book.add_paragraph('法向压力设计值：', style='Normal')
    print('……', end=" ")
    nyplus = -ny
    mathtemp = r'N = ' + str(nyplus) + r'(N)'
    width = add_image(mathtemp, 'ny')
    calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/ny.png', width=Inches(width))
    calc_book.add_paragraph('根据规范，法向压力设计值应满足下式', style='Normal')
    print('……', end=" ")
    fca = 0.5 * fc * mbh * mbw
    mathtemp = r'N \leq 0.5 f_c A =' + str(ceil(fca)) + r'(N)'
    width = add_image(mathtemp, 'fca')
    calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/fca.png', width=Inches(width))
    if nyplus <= fca:
        calc_book.add_paragraph('满足规范要求。', style='Quote').add_run('')
    else:
        calc_book.add_paragraph('', style='Quote').add_run('不满足规范要求。').font.color.rgb = RGBColor(0xff, 0x00, 0x00)

calc_book.add_paragraph('锚筋的抗拉强度设计值：', style='Normal')
print('……', end=" ")
mathtemp = r'f_y = ' + str(fy) + r'(N/mm^2)'
width = add_image(mathtemp, 'fy')
calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/fy.png', width=Inches(width))

calc_book.add_paragraph('混凝土轴心抗压强度设计值：', style='Normal')
print('……', end=" ")
mathtemp = r'f_c = ' + str(fc) + r'(N/mm^2)'
width = add_image(mathtemp, 'fc')
calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/fc.png', width=Inches(width))

calc_book.add_paragraph('混凝土轴心抗拉强度设计值：', style='Normal')
print('……', end=" ")
mathtemp = r'f_t = ' + str(ft) + r'(N/mm^2)'
width = add_image(mathtemp, 'ft')
calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/ft.png', width=Inches(width))

calc_book.add_paragraph('锚筋层数的影响系数：', style='Normal')
print('……', end=" ")
# 插入公式图片开始
mathtemp = r'\alpha_r = ' + str(ar)
width = add_image(mathtemp, 'ar')
calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/ar.png', width=Inches(width))
# 插入公式图片结束

calc_book.add_paragraph('锚筋的受剪承载力系数：', style='Normal')
print('……', end=" ")
# 插入公式图片开始
mathtemp = r'\alpha_v = (4.0 - 0.08d) \sqrt{f_c \over f_y} = ' + str(round(av, 3))
width = add_image(mathtemp, 'av')
calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/av.png', width=Inches(width))
# 插入公式图片结束

calc_book.add_paragraph('锚板的弯曲变形折减系数：', style='Normal')
print('……', end=" ")
mathtemp = r'\alpha_b = 0.6 + 0.25 \frac{t}{d} = ' + str(round(ab, 3))
width = add_image(mathtemp, 'ab')
calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/ab.png', width=Inches(width))

if ny >= 0:  # 受拉工况校核
    if mx < 0.4 * ny * zz:
        calc_book.add_paragraph('根据规范', style='Normal')
        mathtemp = r'M_x =' + str(mx) + r'(N \cdot mm) < 0.4 N z_z =' + str(ceil(0.4 * ny * zz)) + r'(N \cdot mm)'
        width = add_image(mathtemp, 'mnz1')
        calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/mnz1.png', width=Inches(width))
        calc_book.add_paragraph('取', style='Normal')
        mathtemp = r'M_x = 0.4 N z_z =' + str(ceil(0.4 * ny * zz)) + r'(N \cdot mm)'
        width = add_image(mathtemp, 'mnz2')
        calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/mnz2.png', width=Inches(width))
        mx = 0.4 * ny * zz
    if mz < 0.4 * ny * zx:
        calc_book.add_paragraph('根据规范', style='Normal')
        mathtemp = r'M_z =' + str(mz) + r'(N \cdot mm) < 0.4 N z_x =' + str(ceil(0.4 * ny * zx)) + r'(N \cdot mm)'
        width = add_image(mathtemp, 'mnz3')
        calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/mnz3.png', width=Inches(width))
        calc_book.add_paragraph('取', style='Normal')
        mathtemp = r'M_z = 0.4 N z_x =' + str(ceil(0.4 * ny * zx)) + r'(N \cdot mm)'
        width = add_image(mathtemp, 'mnz4')
        calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/mnz4.png', width=Inches(width))
        mz = 0.4 * ny * zx
    as1 = vx / (ar * av * fy) + vz / (ar * av * fy) + ny / (0.8 * ab * fy) + mx / (1.3 * ar * ab * fy * zz) + mz / (
            1.3 * ar * ab * fy * zx)
    as2 = ny / (0.8 * ab * fy) + mx / (0.4 * ar * ab * fy * zz) + mz / (0.4 * ar * ab * fy * zx)
    calc_book.add_paragraph('剪力、法向拉力和弯矩共同作用，锚筋总面积应满足下列两式', style='Normal')
    print('……', end=" ")
    mathtemp = r'A_s \geq \frac{V}{\alpha_r \alpha_v f_y} + \frac{N}{0.8 \alpha_b f_y} ' \
               r'+ \frac{M}{1.3 \alpha_r \alpha_b f_y z}=' + str(ceil(as1)) + r'(mm^2)'
    width = add_image(mathtemp, 'as1')
    calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/as1.png', width=Inches(width))
    print('……', end=" ")
    mathtemp = r'A_s \geq \frac{N}{0.8 \alpha_b f_y} + \frac{M}{0.4 \alpha_r \alpha_b f_y z}=' + str(
        ceil(as2)) + r'(mm^2)'
    width = add_image(mathtemp, 'as2')
    calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/as2.png', width=Inches(width))
    calc_book.add_paragraph('锚筋实际总面积', style='Normal')
    print('……', end=" ")
    mathtemp = r'A_s =' + str(ceil(mjas)) + r'(mm^2)'
    width = add_image(mathtemp, 'as')
    calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/as.png', width=Inches(width))
    if mjas > as1 and mjas > as2:
        calc_book.add_paragraph('满足规范要求。', style='Quote')
    else:
        calc_book.add_paragraph('', style='Quote').add_run('不满足规范要求。').font.color.rgb = RGBColor(0xff, 0x00, 0x00)

if ny < 0:  # 受压工况校核
    ny = -ny  # 临时将负数的压力转为正值，用于计算
    if mx < 0.4 * ny * zz:
        calc_book.add_paragraph('根据规范', style='Normal')
        mathtemp = r'M_x =' + str(mx) + r'(N \cdot mm) < 0.4 N z_z =' + str(ceil(0.4 * ny * zz)) + r'(N \cdot mm)'
        width = add_image(mathtemp, 'mnz1')
        calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/mnz1.png', width=Inches(width))
        calc_book.add_paragraph('取', style='Normal')
        mathtemp = r'M_x = 0.4 N z_z =' + str(ceil(0.4 * ny * zz)) + r'(N \cdot mm)'
        width = add_image(mathtemp, 'mnz2')
        calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/mnz2.png', width=Inches(width))
        mx = 0.4 * ny * zz
    if mz < 0.4 * ny * zx:
        calc_book.add_paragraph('根据规范', style='Normal')
        mathtemp = r'M_z =' + str(mz) + r'(N \cdot mm) < 0.4 N z_x =' + str(ceil(0.4 * ny * zx)) + r'(N \cdot mm)'
        width = add_image(mathtemp, 'mnz3')
        calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/mnz3.png', width=Inches(width))
        calc_book.add_paragraph('取', style='Normal')
        mathtemp = r'M_z = 0.4 N z_x =' + str(ceil(0.4 * ny * zx)) + r'(N \cdot mm)'
        width = add_image(mathtemp, 'mnz4')
        calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/mnz4.png', width=Inches(width))
        mz = 0.4 * ny * zx
    as3 = (vx - 0.3 * ny) / (ar * av * fy) + (vz - 0.3 * ny) / (ar * av * fy) + (mx - 0.4 * ny * zz) / (
            1.3 * ar * ab * fy * zz) + (mz - 0.4 * ny * zx) / (1.3 * ar * ab * fy * zx)
    as4 = (mx - 0.4 * ny * zz) / (0.4 * ar * ab * fy * zz) + (mz - 0.4 * ny * zx) / (0.4 * ar * ab * fy * zx)
    calc_book.add_paragraph('剪力、法向压力和弯矩共同作用，锚筋总面积应满足下列两式', style='Normal')
    mathtemp = r'A_s \geq \frac{V-0.3N}{\alpha_r \alpha_v f_y} + \frac{M-0.4Nz}{1.3 \alpha_r \alpha_b f_y z} =' + str(
        ceil(as3)) + r'(mm^2)'
    width = add_image(mathtemp, 'as3')
    calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/as3.png', width=Inches(width))
    print('……', end=" ")
    mathtemp = r'A_s \geq \frac{M-0.4Nz}{0.4 \alpha_r \alpha_b f_y z} =' + str(ceil(as4)) + r'(mm^2)'
    width = add_image(mathtemp, 'as4')
    calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/as4.png', width=Inches(width))
    calc_book.add_paragraph('锚筋实际总面积', style='Normal')
    print('……', end=" ")
    mathtemp = r'A_s =' + str(ceil(mjas)) + r'(mm^2)'
    width = add_image(mathtemp, 'as')
    calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/as.png', width=Inches(width))
    if mjas > as3 and mjas > as4:
        calc_book.add_paragraph('满足规范要求。', style='Quote')
    else:
        calc_book.add_paragraph('', style='Quote').add_run('不满足规范要求。').font.color.rgb = RGBColor(0xff, 0x00, 0x00)
    ny = -ny  # 恢复压力的负数值

if ny >= 0:  # 受拉工况的锚固长度计算
    calc_book.add_paragraph('钢筋的外形系数：', style='Normal')
    print('……', end=" ")
    mathtemp = r'\alpha = 0.14'
    width = add_image(mathtemp, 'alpha')
    calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/alpha.png', width=Inches(width))
    calc_book.add_paragraph('基本锚固长度：', style='Normal')
    lab = 0.14 * fy * d / ft
    print('……', end=" ")
    mathtemp = r'l_{ab} = \alpha \frac{f_y}{f_t} d = ' + str(ceil(lab)) + '(mm)'
    width = add_image(mathtemp, 'lab')
    calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/lab.png', width=Inches(width))
    calc_book.add_paragraph('受拉钢筋的锚固长度限值：', style='Normal')
    mathtemp = r'l_{a} = \zeta_a l_{ab} = ' + str(ceil(lab * 1.1)) + '(mm)'
    width = add_image(mathtemp, 'la')
    calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/la.png', width=Inches(width))
    calc_book.add_paragraph('上式中，锚固长度修正系数按规范取：', style='Normal')
    mathtemp = r'\zeta_a = ' + str(1.1)
    width = add_image(mathtemp, 'zeta')
    calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/zeta.png', width=Inches(width))
    calc_book.add_paragraph('埋件处墙体厚度不满足锚固长度限值时采取双面锚板加强措施。', style='Normal')

if ny < 0:  # 受压工况的锚固长度计算
    calc_book.add_paragraph('根据规范，受压直锚筋的锚固长度不应小于15d', style='Normal')
    calc_book.add_paragraph('受压钢筋的锚固长度限值：', style='Normal')
    mathtemp = r'l_{a} \geq 15d = ' + str(ceil(15 * d)) + '(mm)'
    width = add_image(mathtemp, 'lay')
    calc_book.add_paragraph('', style='No Spacing').add_run('').add_picture(f'{path}/lay.png', width=Inches(width))

print('……')
'''
计算书结束，输出docx文档
'''
filename = '埋件校核' + strftime("%Y-%m-%d-%H%M%S", localtime())
calc_book.save(f'{filename}.docx')
print(f'计算书生成结束，保存在程序目录下，文件名为{filename}.docx')
input('按回车键退出......')
